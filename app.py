"""
Asistente de Auditoría SGI — v3.0
ISO 9001:2015 / ISO 39001:2015
Stack: Streamlit · ChromaDB · Gemini 2.5 Flash · all-MiniLM-L6-v2 · python-docx
"""

import os, json, hashlib, tempfile, uuid
from datetime import datetime
from pathlib import Path
from io import BytesIO

# ─── CARGAR .env AUTOMÁTICAMENTE ─────────────────────────────────────────────
def _load_dotenv():
    """Carga variables de .env sin dependencias externas."""
    env_path = Path(__file__).parent / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        k = k.strip(); v = v.strip().strip('"').strip("'")
        if k and k not in os.environ:
            os.environ[k] = v
_load_dotenv()

import streamlit as st

# ─── PÁGINA ──────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Asistente SGI · ISO 9001/39001",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CONSTANTES ──────────────────────────────────────────────────────────────
LISTA_MAESTRA_PATH   = Path("lista_maestra.json")
ANALISIS_PATH        = Path("analisis_documentos.json")
REVISIONES_PATH      = Path("revisiones.json")
SUGERENCIAS_NUE_PATH = Path("sugerencias_nuevos.json")
INCONGRUENCIAS_PATH  = Path("incongruencias.json")
AGENDA_PATH          = Path("agenda_sgi.json")
HALLAZGOS_PATH       = Path("hallazgos_auditoria.json")
CHROMA_PATH          = "./chroma_db"
CHROMA_COLLECTION    = "sgi_documentos"

CHUNK_SIZE       = 700
CHUNK_OVERLAP    = 80
RAG_TOP_K        = 3
TEXT_PREVIEW_LEN = 6000

# ─── CSS ─────────────────────────────────────────────────────────────────────
CSS = """
<style>
[data-testid="stAppViewContainer"] { background:#f0f4f8 !important; }
[data-testid="stMain"] { padding-top:0 !important; }
section[data-testid="stSidebar"] {
    background:linear-gradient(180deg,#0d47a1 0%,#0f3460 100%) !important;
    border-right:none !important;
}
section[data-testid="stSidebar"] * { color:rgba(255,255,255,.9) !important; }
section[data-testid="stSidebar"] hr { border-color:rgba(255,255,255,.15) !important; }
section[data-testid="stSidebar"] h3 { color:#90caf9 !important; font-size:.8rem !important;
    text-transform:uppercase; letter-spacing:.7px; }

.sgi-header {
    background:linear-gradient(135deg,#0d47a1 0%,#1565c0 45%,#0288d1 100%);
    padding:26px 30px 22px; border-radius:16px; margin-bottom:18px;
    box-shadow:0 6px 28px rgba(13,71,161,.45); position:relative; overflow:hidden;
}
.sgi-header::before { content:''; position:absolute; right:-40px; top:-40px;
    width:220px; height:220px; border-radius:50%; background:rgba(255,255,255,.06); }
.sgi-header h1 { color:#fff; margin:0; font-size:1.9rem; font-weight:800; letter-spacing:-.5px; }
.sgi-header p  { color:#bbdefb; margin:7px 0 0; font-size:.95rem; }
.kpi-row { display:flex; gap:10px; margin-top:16px; flex-wrap:wrap; }
.kpi-badge { background:rgba(255,255,255,.14); backdrop-filter:blur(6px);
    border:1px solid rgba(255,255,255,.22); border-radius:12px; padding:9px 16px;
    color:white; min-width:90px; }
.kpi-badge .kv { font-size:1.45rem; font-weight:800; display:block; line-height:1.1; }
.kpi-badge .kl { font-size:.68rem; opacity:.78; text-transform:uppercase; letter-spacing:.6px; }

.sec-title { display:flex; align-items:center; gap:8px; font-size:.98rem; font-weight:700;
    color:#1a237e; border-left:4px solid #1565c0; padding-left:10px; margin:0 0 14px 0; }

.card { background:#fff; border-radius:14px; padding:18px 20px;
    box-shadow:0 2px 14px rgba(0,0,0,.07); margin-bottom:12px;
    border:1px solid rgba(0,0,0,.05); color:#1e293b !important; }
.card-warn { background:#fffbf0; border-left:4px solid #f59e0b;
    padding:11px 15px; border-radius:0 12px 12px 0; margin-bottom:10px;
    box-shadow:0 2px 8px rgba(245,158,11,.12); color:#1e293b !important; }
.card-ok   { background:#f0fdf4; border-left:4px solid #22c55e;
    padding:11px 15px; border-radius:0 12px 12px 0; margin-bottom:10px;
    box-shadow:0 2px 8px rgba(34,197,94,.12); color:#1e293b !important; }
.card-info { background:#eff6ff; border-left:4px solid #3b82f6;
    padding:11px 15px; border-radius:0 12px 12px 0; margin-bottom:10px;
    box-shadow:0 2px 8px rgba(59,130,246,.12); color:#1e293b !important; }
.card-danger { background:#fff1f2; border-left:4px solid #ef4444;
    padding:11px 15px; border-radius:0 12px 12px 0; margin-bottom:10px;
    box-shadow:0 2px 8px rgba(239,68,68,.12); color:#1e293b !important; }

.verdict { display:inline-flex; align-items:center; gap:6px; padding:7px 18px;
    border-radius:999px; font-weight:700; font-size:1rem; letter-spacing:.3px; margin-bottom:12px; }
.verdict-ok  { background:#dcfce7; color:#15803d; border:2px solid #22c55e; }
.verdict-obs { background:#fef9c3; color:#a16207; border:2px solid #eab308; }
.verdict-nok { background:#fee2e2; color:#b91c1c; border:2px solid #ef4444; }

.doc-badge  { display:inline-block; padding:3px 10px; border-radius:999px;
    font-size:.71rem; font-weight:600; letter-spacing:.3px; }
.badge-proc  { background:#e0f2fe; color:#0369a1; }
.badge-reg   { background:#f0fdf4; color:#15803d; }
.badge-man   { background:#faf5ff; color:#6b21a8; }
.badge-pol   { background:#fff7ed; color:#c2410c; }
.badge-inst  { background:#f0fdfa; color:#0f766e; }
.badge-inf   { background:#fdf4ff; color:#86198f; }
.badge-other { background:#f1f5f9; color:#475569; }

.status-pill { display:inline-block; padding:2px 10px; border-radius:999px;
    font-size:.71rem; font-weight:700; letter-spacing:.3px; }
.st-abierta  { background:#fef3c7; color:#92400e; }
.st-resuelta { background:#d1fae5; color:#065f46; }
.st-eliminada{ background:#f3f4f6; color:#9ca3af; }

.inc-card { background:#fff; border-radius:12px; padding:14px 16px;
    box-shadow:0 2px 10px rgba(0,0,0,.06); margin-bottom:8px;
    border:1.5px solid #e5e7eb; transition:border-color .2s; color:#1e293b !important; }
.inc-card-active { border-color:#1565c0 !important; background:#eff6ff !important; }

.repo-card { background:#fff; border-radius:14px; padding:16px 18px;
    box-shadow:0 2px 10px rgba(0,0,0,.06); margin-bottom:10px;
    border:1px solid #e5e7eb; color:#1e293b !important; }
.obsoleto-banner { background:linear-gradient(90deg,#7f1d1d,#991b1b);
    color:white !important; font-weight:700; font-size:.78rem; letter-spacing:1px;
    text-transform:uppercase; padding:4px 12px; border-radius:6px;
    display:inline-block; margin-bottom:8px; }

.stat-card { background:rgba(255,255,255,.1); border-radius:12px; padding:12px;
    margin-bottom:8px; text-align:center; border:1px solid rgba(255,255,255,.18); }
.stat-val { font-size:1.6rem; font-weight:800; color:white; line-height:1; }
.stat-lbl { font-size:.68rem; color:rgba(255,255,255,.7); text-transform:uppercase; letter-spacing:.5px; }

.step-item { display:flex; gap:12px; align-items:flex-start; padding:10px 14px;
    border-radius:10px; margin-bottom:8px; background:#f8fafc; border:1px solid #e2e8f0; }
.step-num { background:#1565c0; color:white; border-radius:50%; width:24px; height:24px;
    display:flex; align-items:center; justify-content:center; font-size:.75rem; font-weight:700;
    flex-shrink:0; margin-top:1px; }
.step-txt { font-size:.87rem; color:#374151; line-height:1.5; }

div[data-testid="stTabs"] [role="tablist"] { background:#e3eaf7; border-radius:14px; padding:5px; gap:4px; }
div[data-testid="stTabs"] [role="tab"] { border-radius:10px !important; font-weight:600 !important;
    font-size:.85rem !important; padding:8px 18px !important; color:#4a5568 !important;
    transition:all .2s ease !important; border:none !important; }
div[data-testid="stTabs"] [role="tab"][aria-selected="true"] { background:#1565c0 !important;
    color:white !important; box-shadow:0 3px 10px rgba(21,101,192,.4) !important; }
div[data-testid="stTabs"] [role="tab"]:hover:not([aria-selected="true"]) {
    background:rgba(21,101,192,.1) !important; color:#1565c0 !important; }

.rpanel-lbl { font-size:.75rem; font-weight:600; text-transform:uppercase;
    letter-spacing:.6px; color:#6b7280; margin-bottom:6px; }
.foot-info { font-size:.72rem; color:rgba(255,255,255,.6); line-height:1.9; text-align:center; }
div[data-testid="stChatMessage"] { border-radius:12px; padding:4px 2px; }
div[data-testid="stChatMessage"] p,
div[data-testid="stChatMessage"] li,
div[data-testid="stChatMessage"] span,
div[data-testid="stChatMessage"] div { color:#1e293b !important; }
div[data-testid="stChatMessageContent"] { color:#1e293b !important; }
hr { border-color:#e2e8f0 !important; }

/* ── Texto visible en alertas (error / warning / info / success) ────────── */
div[data-testid="stAlert"] p,
div[data-testid="stAlert"] li,
div[data-testid="stAlert"] span,
div[data-testid="stAlert"] a { color:#1e293b !important; }
div[data-testid="stNotification"] * { color:#1e293b !important; }
.stAlert > div { color:#1e293b !important; }

/* ── Scrollbar visible en contenedores de chat ─────────────────────────── */
div[data-testid="stVerticalBlockBorderWrapper"] {
    scrollbar-width: thin !important;
    scrollbar-color: #1565c0 #e3eaf7 !important;
}
div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar {
    width: 7px !important;
}
div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar-track {
    background: #e3eaf7 !important;
    border-radius: 4px !important;
}
div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar-thumb {
    background: #1565c0 !important;
    border-radius: 4px !important;
    border: 1px solid #e3eaf7 !important;
}
div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar-thumb:hover {
    background: #0d47a1 !important;
}

/* Lista Maestra / Matriz */
.matriz-table { width:100%; border-collapse:collapse; font-size:.82rem; }
.matriz-table th { background:#1565c0; color:white; padding:9px 12px; text-align:left;
    font-weight:700; font-size:.76rem; letter-spacing:.4px; }
.matriz-table tr:nth-child(even) { background:#f8fafc; }
.matriz-table tr:hover { background:#eff6ff; }
.matriz-table td { padding:8px 12px; border-bottom:1px solid #e2e8f0; vertical-align:top;
    color:#1e293b; }
.clausula-tag { display:inline-block; background:#dbeafe; color:#1d4ed8; border-radius:6px;
    padding:2px 7px; font-size:.7rem; font-weight:600; margin:2px 2px 2px 0; }
.clausula-tag-39 { background:#f3e8ff; color:#7c3aed; }

/* Agenda SGI */
.agenda-card { background:#fff; border-radius:14px; padding:16px 18px; margin-bottom:10px;
    border:1.5px solid #e5e7eb; box-shadow:0 2px 8px rgba(0,0,0,.05); color:#1e293b !important; }
.agenda-card-done { border-color:#22c55e !important; background:#f0fdf4 !important; }
.agenda-card-vencido { border-color:#ef4444 !important; background:#fff1f2 !important; }
.freq-pill { display:inline-block; background:#e0f2fe; color:#0369a1; border-radius:999px;
    padding:2px 10px; font-size:.69rem; font-weight:700; letter-spacing:.3px; }
.freq-pill-mes { background:#fef9c3; color:#a16207; }
.freq-pill-trim { background:#dcfce7; color:#15803d; }
.freq-pill-anual { background:#f3e8ff; color:#6b21a8; }

/* Hallazgos */
.hallazgo-card { background:#fff; border-radius:14px; padding:16px 18px; margin-bottom:10px;
    border-left:5px solid #e5e7eb; box-shadow:0 2px 8px rgba(0,0,0,.05); color:#1e293b !important; }
.hallazgo-nc-mayor { border-left-color:#ef4444 !important; }
.hallazgo-nc-menor { border-left-color:#f59e0b !important; }
.hallazgo-obs      { border-left-color:#3b82f6 !important; }
.hallazgo-opm      { border-left-color:#22c55e !important; }
.htipo-pill { display:inline-block; padding:3px 10px; border-radius:999px; font-size:.71rem;
    font-weight:700; letter-spacing:.3px; }
.htipo-ncm  { background:#fee2e2; color:#b91c1c; }
.htipo-ncn  { background:#fef3c7; color:#92400e; }
.htipo-obs  { background:#dbeafe; color:#1e40af; }
.htipo-opm  { background:#d1fae5; color:#065f46; }
.hestado-pill { display:inline-block; padding:2px 9px; border-radius:999px; font-size:.69rem; font-weight:700; }
.he-abierto    { background:#fee2e2; color:#b91c1c; }
.he-en_proceso { background:#fef9c3; color:#a16207; }
.he-cerrado    { background:#d1fae5; color:#065f46; }
.he-verificado { background:#dbeafe; color:#1e40af; }
</style>
"""

# ─── JSON HELPERS ─────────────────────────────────────────────────────────────
def load_json(path: Path, default):
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        pass
    return default

def save_json(path: Path, data) -> bool:
    try:
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return True
    except Exception as e:
        st.error(f"Error guardando {path.name}: {e}")
        return False

# ─── CHROMADB + EMBEDDINGS ────────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def _init_chroma():
    import chromadb
    from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
    ef = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    client = chromadb.PersistentClient(path=CHROMA_PATH)
    return client.get_or_create_collection(
        name=CHROMA_COLLECTION, embedding_function=ef,
        metadata={"hnsw:space": "cosine"},
    )

def get_collection():
    return _init_chroma()

def get_document_text_from_chroma(filename: str) -> str:
    try:
        res = get_collection().get(where={"source": filename}, include=["documents"])
        chunks = res.get("documents", [])
        return " ".join(chunks)[:TEXT_PREVIEW_LEN] if chunks else ""
    except Exception:
        return ""

# ─── OCR CON GEMINI VISION ───────────────────────────────────────────────────
def ocr_with_gemini(file_bytes: bytes, filename: str) -> str:
    """OCR de documentos escaneados usando Gemini Files API."""
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key:
        return ""
    if len(file_bytes) > 50 * 1024 * 1024:
        st.warning("⚠️ El PDF supera 50 MB. Usá ilovepdf.com → OCR y volvé a subirlo.")
        return ""
    try:
        from google import genai
        # OCR: intentamos con v1beta primero (2.0), luego v1 (1.5)
        client = genai.Client(api_key=api_key, http_options={"api_version": "v1beta"})
        prompt_ocr = (
            "Extraé TODO el texto de este documento escaneado, página por página. "
            "Devolvé únicamente el texto extraído manteniendo párrafos y secciones separadas. "
            "No incluyas explicaciones ni comentarios propios."
        )
        # ── Método 1: Files API (más robusto para PDFs) ──────────────────────
        try:
            suffix = Path(filename).suffix.lower() or ".pdf"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                uploaded = client.files.upload(file=tmp_path)
                for model_id in ("gemini-2.0-flash", "gemini-2.0-flash-lite",
                                  "gemini-1.5-flash-latest", "gemini-1.5-flash-001"):
                    try:
                        resp = client.models.generate_content(
                            model=model_id,
                            contents=[uploaded, prompt_ocr],
                        )
                        txt = (resp.text or "").strip()
                        if txt:
                            return txt
                    except Exception:
                        continue
            finally:
                os.unlink(tmp_path)
        except Exception:
            pass
        # ── Método 2: inline bytes (fallback) ────────────────────────────────
        try:
            from google.genai import types as gtypes
            ocr_models = [
                ("gemini-2.0-flash",  "v1beta"),
                ("gemini-1.5-flash",  "v1"),
                ("gemini-1.5-pro",    "v1"),
            ]
            for model_id, api_ver in ocr_models:
                try:
                    c2 = genai.Client(api_key=api_key,
                                      http_options={"api_version": api_ver})
                    resp = c2.models.generate_content(
                        model=model_id,
                        contents=[
                            gtypes.Part.from_bytes(
                                data=file_bytes, mime_type="application/pdf"
                            ),
                            prompt_ocr,
                        ],
                    )
                    txt = (resp.text or "").strip()
                    if txt:
                        return txt
                except Exception:
                    continue
        except Exception:
            pass
    except Exception as e:
        st.warning(f"OCR Gemini error: {e}")
    return ""


# ─── EXTRACCIÓN DE TEXTO ──────────────────────────────────────────────────────
def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
            # ── Detección de PDF escaneado (sin texto extraíble) ──
            if len(text.strip()) < 120:
                st.info("📷 PDF escaneado detectado — aplicando OCR con Gemini Vision…")
                ocr_text = ocr_with_gemini(file_bytes, filename)
                if ocr_text:
                    st.success("✅ OCR completado correctamente.")
                    return ocr_text
                else:
                    st.error(
                        "❌ No se pudo extraer texto. Opciones:\n"
                        "1. **ilovepdf.com** → OCR → descargá el PDF con texto → volvé a subirlo\n"
                        "2. **Adobe Acrobat** → Herramientas → Reconocer texto (OCR)\n"
                        "3. **Microsoft Lens** → escanear con teléfono → exportar como Word"
                    )
                    return ""
            return text
        if ext == ".docx":
            import docx2txt
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(file_bytes); tmp_path = tmp.name
            try:    return docx2txt.process(tmp_path) or ""
            finally: os.unlink(tmp_path)
        if ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"\n=== Hoja: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    r = " | ".join(str(c) if c is not None else "" for c in row)
                    if r.replace("|","").strip(): lines.append(r)
            return "\n".join(lines)
        if ext in (".html", ".htm"):
            import re, html as _html
            raw = file_bytes.decode("utf-8", errors="ignore")
            # Quitar scripts y estilos primero
            raw = re.sub(r'<(script|style)[^>]*>.*?</(script|style)>', ' ', raw,
                         flags=re.IGNORECASE | re.DOTALL)
            # Quitar etiquetas HTML
            text = re.sub(r'<[^>]+>', ' ', raw)
            text = _html.unescape(text)
            return re.sub(r'\s+', ' ', text).strip()
        if ext == ".txt":
            return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Error extrayendo '{filename}': {e}")
    return ""

def chunk_text(text: str) -> list[str]:
    words = text.split()
    if not words: return []
    step = max(1, CHUNK_SIZE - CHUNK_OVERLAP)
    chunks = []
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i+CHUNK_SIZE])
        if chunk.strip(): chunks.append(chunk)
        if i + CHUNK_SIZE >= len(words): break
    return chunks

# ─── GEMINI ───────────────────────────────────────────────────────────────────
def _gemini_client():
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key: return None
    try:
        from google import genai
        return genai.Client(api_key=api_key)
    except ImportError:
        return None

# ── GROQ FALLBACK ─────────────────────────────────────────────────────────────
def _call_groq(prompt: str) -> str | None:
    """Fallback gratuito usando Groq (Llama 3.3 70B). Requiere GROQ_API_KEY en .env"""
    api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return None
    try:
        from groq import Groq
        client = Groq(api_key=api_key)
        # Modelos en orden de preferencia
        for model in ("llama-3.3-70b-versatile", "llama-3.1-70b-versatile",
                      "mixtral-8x7b-32768", "llama-3.1-8b-instant"):
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=4096,
                    temperature=0.2,
                )
                txt = (resp.choices[0].message.content or "").strip()
                if txt:
                    return txt
            except Exception as e:
                if "model_not_found" in str(e).lower() or "not found" in str(e).lower():
                    continue  # prueba el siguiente modelo
                raise
    except Exception:
        pass
    return None

def _call_gemini(prompt: str) -> str | None:
    """Llama a Gemini 2.0 Flash. Si la cuota se agota, usa Groq automáticamente."""
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key:
        # Sin Gemini → intentar directo con Groq
        return _call_groq(prompt)

    try:
        from google import genai
    except ImportError:
        st.error("Instale google-genai: pip install google-genai")
        return None

    # gemini-2.0-flash-lite tiene 30 RPM (el doble que flash), va primero
    MODELS = ["gemini-2.0-flash-lite", "gemini-2.0-flash"]
    client = genai.Client(api_key=api_key)

    import time
    quota_hit = False
    for model_id in MODELS:
        try:
            resp = client.models.generate_content(model=model_id, contents=prompt)
            txt = (resp.text or "").strip()
            if txt:
                return txt
        except Exception as e:
            err_str = str(e)
            if "RESOURCE_EXHAUSTED" in err_str or "quota" in err_str.lower():
                quota_hit = True
                time.sleep(3)   # pequeña pausa antes de probar el siguiente modelo
                continue
            continue

    # Gemini falló — intentar con Groq
    groq_key = os.environ.get("GROQ_API_KEY", "")
    if groq_key:
        result = _call_groq(prompt)
        if result:
            return result   # ← Groq respondió sin mostrar ningún aviso al usuario

    # Todo falló
    if quota_hit:
        if groq_key:
            st.warning("⏳ Cuota de Gemini agotada y Groq tampoco respondió. "
                       "Esperá 1 minuto y reintentá.")
        else:
            st.warning("⏳ Cuota de Gemini agotada. **Opciones:**\n"
                       "1. Esperá 1-2 minutos y reintentá\n"
                       "2. Agregá `GROQ_API_KEY` en el archivo `.env` para usar "
                       "Llama 3.3 como respaldo gratuito → [groq.com](https://console.groq.com/keys)")
    return None

def _parse_json(raw: str | None) -> dict | None:
    if not raw: return None
    if "```" in raw:
        for block in raw.split("```")[1::2]:
            try: return json.loads(block.lstrip("json").strip())
            except Exception: continue
    try: return json.loads(raw)
    except Exception: return None

# ─── DICCIONARIOS ISO (para Matriz Normativa) ────────────────────────────────
ISO_9001_DICT = {
    "4.1":"Contexto de la organización",
    "4.2":"Partes interesadas",
    "4.3":"Alcance del SGC",
    "4.4":"SGC y sus procesos",
    "5.1":"Liderazgo y compromiso",
    "5.2":"Política de la calidad",
    "5.3":"Roles y responsabilidades",
    "6.1":"Riesgos y oportunidades",
    "6.2":"Objetivos de calidad",
    "6.3":"Planificación de cambios",
    "7.1":"Recursos",
    "7.2":"Competencia",
    "7.3":"Toma de conciencia",
    "7.4":"Comunicación",
    "7.5":"Información documentada",
    "8.1":"Planificación operacional",
    "8.2":"Requisitos de productos/servicios",
    "8.3":"Diseño y desarrollo",
    "8.4":"Proveedores externos",
    "8.5":"Producción y prestación",
    "8.6":"Liberación de productos",
    "8.7":"Salidas no conformes",
    "9.1":"Seguimiento y medición",
    "9.2":"Auditoría interna",
    "9.3":"Revisión por la dirección",
    "10.2":"No conformidad y AC",
    "10.3":"Mejora continua",
}
ISO_39001_DICT = {
    "4.1":"Contexto de la organización",
    "4.2":"Partes interesadas",
    "4.3":"Alcance del SGSV",
    "4.4":"SGSV",
    "5.1":"Liderazgo y compromiso",
    "5.2":"Política de seguridad vial",
    "5.3":"Roles y responsabilidades",
    "6.1":"Factores de riesgo viales",
    "6.2":"Metas de comportamiento vial",
    "7.1":"Recursos",
    "7.2":"Competencia",
    "7.3":"Toma de conciencia",
    "7.4":"Comunicación",
    "7.5":"Información documentada",
    "8.1":"Planificación operacional",
    "8.3":"Gestión de riesgos viales",
    "8.4":"Indicadores de desempeño SV",
    "9.1":"Seguimiento y medición",
    "9.2":"Auditoría interna",
    "9.3":"Revisión por la dirección",
    "10.1":"No conformidad y AC",
    "10.2":"Mejora continua",
}

# ─── ACTIVIDADES OBLIGATORIAS DEL SGI ────────────────────────────────────────
AGENDA_DEFAULT = [
    {"id":"A001","clausula_9001":"4.1 / 4.2","clausula_39001":"4.1 / 4.2",
     "nombre":"Análisis de Contexto y Partes Interesadas",
     "descripcion":"Revisar y actualizar el análisis de factores internos/externos, la matriz de partes interesadas y sus necesidades. Actualizar FODA/PESTEL.",
     "frecuencia":"Anual","responsable":"Alta Dirección","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A002","clausula_9001":"5.1","clausula_39001":"5.1",
     "nombre":"Evidencia de Compromiso de Liderazgo",
     "descripcion":"Registrar evidencias del compromiso de la alta dirección: actas de reuniones, comunicaciones, asignación de recursos al SGI.",
     "frecuencia":"Semestral","responsable":"Alta Dirección","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A003","clausula_9001":"6.1","clausula_39001":"6.1",
     "nombre":"Evaluación de Riesgos y Oportunidades",
     "descripcion":"Actualizar la matriz de riesgos y oportunidades. Verificar eficacia de las acciones implementadas anteriormente.",
     "frecuencia":"Semestral","responsable":"Responsable SGI","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A004","clausula_9001":"6.2","clausula_39001":"6.2",
     "nombre":"Revisión de Objetivos de Calidad y SV",
     "descripcion":"Evaluar el grado de cumplimiento de los objetivos vigentes con sus indicadores. Establecer nuevos objetivos si corresponde.",
     "frecuencia":"Semestral","responsable":"Responsable SGI","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A005","clausula_9001":"7.2","clausula_39001":"7.2",
     "nombre":"Evaluación de Competencias del Personal",
     "descripcion":"Revisar perfiles de puestos, registros de capacitación y evaluaciones de competencias. Detectar necesidades de formación.",
     "frecuencia":"Anual","responsable":"RRHH / Responsable SGI","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A006","clausula_9001":"8.1","clausula_39001":"8.1",
     "nombre":"Revisión de Planificación Operacional",
     "descripcion":"Verificar que los procesos operacionales se ejecutan según lo planificado. Revisar indicadores de proceso y verificar cumplimiento de criterios.",
     "frecuencia":"Trimestral","responsable":"Jefes de Área","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A007","clausula_9001":"9.1","clausula_39001":"9.1",
     "nombre":"Seguimiento y Medición de Indicadores SGI",
     "descripcion":"Recopilar, analizar y reportar los indicadores de desempeño del SGI. Preparar informe de seguimiento para la dirección.",
     "frecuencia":"Trimestral","responsable":"Responsable SGI","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A008","clausula_9001":"9.2","clausula_39001":"9.2",
     "nombre":"Programa de Auditoría Interna",
     "descripcion":"Planificar, ejecutar y documentar las auditorías internas del SGI según el programa anual. Cubrir todas las cláusulas aplicables.",
     "frecuencia":"Anual","responsable":"Auditor Interno","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A009","clausula_9001":"9.3","clausula_39001":"9.3",
     "nombre":"Revisión por la Dirección",
     "descripcion":"Realizar la reunión formal de revisión por la dirección. Puntos obligatorios: estado de objetivos, resultados de auditorías, NC abiertas, satisfacción de partes interesadas, recursos.",
     "frecuencia":"Anual","responsable":"Alta Dirección","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A010","clausula_9001":"10.2","clausula_39001":"10.1",
     "nombre":"Seguimiento de No Conformidades y Acciones Correctivas",
     "descripcion":"Revisar el estado de todas las NC abiertas. Verificar eficacia de las acciones correctivas. Actualizar el registro de AC.",
     "frecuencia":"Mensual","responsable":"Responsable SGI","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A011","clausula_9001":"N/A","clausula_39001":"6.1",
     "nombre":"Relevamiento de Factores de Riesgo Viales",
     "descripcion":"Identificar y actualizar los factores de riesgo viales: velocidad, infraestructura, condición del vehículo, conducta del conductor, entorno vial.",
     "frecuencia":"Anual","responsable":"Responsable de SV Vial","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
    {"id":"A012","clausula_9001":"N/A","clausula_39001":"8.4",
     "nombre":"Medición de Indicadores de Seguridad Vial",
     "descripcion":"Actualizar indicadores SV: accidentes/incidentes viales, km recorridos, horas de capacitación en conducción segura, infracciones.",
     "frecuencia":"Trimestral","responsable":"Responsable de SV Vial","estado":"pendiente",
     "fecha_limite":"","fecha_realizacion":"","observaciones":"","ai_feedback":""},
]

def load_agenda() -> list:
    data = load_json(AGENDA_PATH, None)
    if data is None:
        save_json(AGENDA_PATH, AGENDA_DEFAULT)
        return [dict(a) for a in AGENDA_DEFAULT]
    # Migración: agregar actividades nuevas que no existan
    existing_ids = {a["id"] for a in data}
    added = False
    for act in AGENDA_DEFAULT:
        if act["id"] not in existing_ids:
            data.append(dict(act)); added = True
    if added: save_json(AGENDA_PATH, data)
    return data

def load_hallazgos() -> list:
    return load_json(HALLAZGOS_PATH, [])

def save_hallazgos(items: list):
    save_json(HALLAZGOS_PATH, items)

# ─── FUNCIONES DE ANÁLISIS IA ─────────────────────────────────────────────────
ISO_9001_CLAUSULAS = """ISO 9001:2015 — cláusulas clave:
4.1 Contexto org · 4.2 Partes interesadas · 4.4 SGC y sus procesos
5.1 Liderazgo · 6.1 Riesgos y oportunidades · 6.2 Objetivos de calidad
7.2 Competencia · 7.5 Información documentada (versión, aprobación, distribución)
8.1 Planificación operacional · 8.2 Requisitos de productos/servicios
9.1 Seguimiento y medición · 9.2 Auditoría interna · 9.3 Revisión por dirección
10.2 No conformidad y acción correctiva"""

ISO_39001_CLAUSULAS = """ISO 39001:2015 — cláusulas clave:
4.1 Factores internos/externos · 5.2 Política de SV vial
6.1 Factores de riesgo viales · 6.2 Metas de comportamiento vial
8.1 Planificación operacional · 8.3 Gestión de riesgos viales
8.4 Factores de desempeño en SV · 9.1 Seguimiento y medición
10. Mejora continua"""

def classify_document(text: str, filename: str) -> dict:
    msg = "Configure GEMINI_API_KEY" if not os.environ.get("GEMINI_API_KEY") else "Error"
    prompt = f"""Eres auditor certificado ISO 9001:2015 e ISO 39001:2015.
Analiza "{filename}" y clasifícalo. Responde SOLO JSON válido:
{{
  "tipo": "<Procedimiento|Instructivo|Registro|Manual|Política|Plan|Informe|Formulario|Otro>",
  "iso9001":  "<Capítulos ISO 9001:2015 aplicables con nombre, o 'No aplica'>",
  "iso39001": "<Capítulos ISO 39001:2015 aplicables con nombre, o 'No aplica'>"
}}
TEXTO (primeros 2500 caracteres):
{text[:2500]}"""
    r = _parse_json(_call_gemini(prompt))
    return r if r else {"tipo":"Sin clasificar","iso9001":msg,"iso39001":msg}


def analyze_document_deep(text: str, filename: str) -> dict:
    if not os.environ.get("GEMINI_API_KEY"):
        return {"incongruencias":["⚠️ GEMINI_API_KEY no configurada."],
                "sugerencias":["Configure la clave y vuelva a subir el archivo."]}
    prompt = f"""Eres Auditor Senior de SGI.

{ISO_9001_CLAUSULAS}

{ISO_39001_CLAUSULAS}

Analiza "{filename}" verificando sistemáticamente las cláusulas relevantes.
Responde SOLO JSON válido (4 ítems por lista):
{{
  "incongruencias": [
    "Cláusula X.X: descripción del vacío/no conformidad detectada",
    "Cláusula Y.Y: ...", "...", "..."
  ],
  "sugerencias": [
    "Acción concreta o KPI para resolver incongruencia 1",
    "...", "...", "..."
  ]
}}
TEXTO del documento (primeros 5000 caracteres):
{text[:5000]}"""
    r = _parse_json(_call_gemini(prompt))
    if r and "incongruencias" in r: return r
    return {"incongruencias":["Error al procesar."],"sugerencias":["Reintente."]}


def analyze_revision(old_text: str, new_text: str, old_name: str, new_name: str) -> dict:
    """Comparación real: documento anterior vs nuevo vs cláusulas ISO."""
    prompt = f"""Eres Auditor Senior de SGI experto en gestión de cambios documentales.

{ISO_9001_CLAUSULAS}

{ISO_39001_CLAUSULAS}

Realiza una COMPARACIÓN SISTEMÁTICA entre la versión anterior y la nueva del documento.
Verifica cláusula por cláusula si los cambios son normativamente adecuados.

DOCUMENTO ORIGINAL: {old_name}
VERSIÓN NUEVA:      {new_name}

VERSIÓN ANTERIOR (4000 caracteres):
{old_text[:4000]}

NUEVA VERSIÓN (4000 caracteres):
{new_text[:4000]}

Responde SOLO JSON válido:
{{
  "veredicto": "<APROBADO | CON OBSERVACIONES | RECHAZADO>",
  "resumen_ejecutivo": "<2-3 oraciones del cambio principal>",
  "clausulas_verificadas": [
    {{"clausula":"ISO 9001 7.5","resultado":"OK","detalle":"La nueva versión incluye código y fecha de revisión"}},
    {{"clausula":"ISO 9001 9.2","resultado":"OBSERVACION","detalle":"No especifica frecuencia de auditorías"}},
    {{"clausula":"ISO 39001 8.3","resultado":"NO_APLICA","detalle":"Documento no relacionado con seguridad vial"}}
  ],
  "cambios_detectados": ["Cambio 1","Cambio 2","Cambio 3"],
  "riesgos": ["Riesgo 1","Riesgo 2"],
  "recomendaciones": ["Recomendación 1","Recomendación 2"]
}}"""
    r = _parse_json(_call_gemini(prompt))
    if r and "veredicto" in r: return r
    return {"veredicto":"ERROR","resumen_ejecutivo":"Error al procesar.",
            "clausulas_verificadas":[],"cambios_detectados":[],"riesgos":[],"recomendaciones":[]}


def suggest_placement(new_text: str, new_name: str, lista_maestra: list) -> dict:
    docs_summary = "\n".join(
        f"  - {d['nombre']} | {d['tipo']} | 9001:{d['iso9001'][:50]} | 39001:{d['iso39001'][:50]}"
        for d in lista_maestra[:25]
    ) or "  (Sin documentos aún)"
    prompt = f"""Eres Consultor SGI experto en diseño de sistemas documentales.

DOCUMENTOS ACTUALES DEL SGI:
{docs_summary}

NUEVO DOCUMENTO: {new_name}
CONTENIDO (3500 caracteres):
{new_text[:3500]}

Responde SOLO JSON válido:
{{
  "tipo_sugerido": "<tipo>",
  "iso9001_sugerido": "<cláusulas aplicables>",
  "iso39001_sugerido": "<cláusulas aplicables o 'No aplica'>",
  "ubicacion_sgi": "<descripción de dónde encaja en el SGI>",
  "documentos_relacionados": ["doc1.pdf","doc2.pdf"],
  "vacios_que_cubre": ["vacío 1","vacío 2"],
  "razon": "<justificación>"
}}"""
    r = _parse_json(_call_gemini(prompt))
    if r and "tipo_sugerido" in r: return r
    return {"tipo_sugerido":"Error","iso9001_sugerido":"N/A","iso39001_sugerido":"N/A",
            "ubicacion_sgi":"Error","documentos_relacionados":[],"vacios_que_cubre":[],"razon":"Reintente."}


def chat_incongruencia(inc_texto: str, doc_nombre: str, historial: list, pregunta: str, analisis_cache: dict) -> str:
    """Chat contextual para resolver una incongruencia específica."""
    doc_texto = analisis_cache.get(doc_nombre, {}).get("_texto", "") or get_document_text_from_chroma(doc_nombre)
    conv = "\n".join(
        f"{'Auditor' if m['role']=='user' else 'Asistente'}: {m['content']}"
        for m in historial[-6:]
    )
    prompt = f"""Eres Auditor Senior de SGI experto en ISO 9001:2015 e ISO 39001:2015.
Estás ayudando a RESOLVER una incongruencia específica detectada en el documento "{doc_nombre}".

INCONGRUENCIA A RESOLVER:
{inc_texto}

EXTRACTO DEL DOCUMENTO (contexto):
{doc_texto[:3000]}

{ISO_9001_CLAUSULAS}

HISTORIAL DE CONVERSACIÓN:
{conv}

NUEVA PREGUNTA DEL AUDITOR: {pregunta}

Proporciona orientación concreta, práctica y específica. Si el usuario pide texto para corregir el documento, genera el texto exacto que debería insertarse."""
    return _call_gemini(prompt) or "Error al generar respuesta. Verifique la API Key."


def analyze_sgi_activity(actividad: dict, docs_en_sgi: list) -> str:
    """Analiza el estado de una actividad del SGI y da retroalimentación con documentos relacionados."""
    docs_txt = "\n".join(
        f"  - {d['nombre']} | {d['tipo']} | 9001:{d.get('iso9001','')[:60]}"
        for d in docs_en_sgi[:20]
    ) or "  (Sin documentos cargados)"
    obs = actividad.get("observaciones","") or "(sin observaciones ingresadas)"
    prompt = f"""Eres un consultor experto en SGI certificado en ISO 9001:2015 e ISO 39001:2015.

ACTIVIDAD DEL SGI A EVALUAR:
- Nombre: {actividad['nombre']}
- Cláusula ISO 9001: {actividad['clausula_9001']}
- Cláusula ISO 39001: {actividad['clausula_39001']}
- Descripción: {actividad['descripcion']}
- Frecuencia requerida: {actividad['frecuencia']}
- Estado actual: {actividad['estado']}
- Fecha realización: {actividad.get('fecha_realizacion','No registrada')}
- Observaciones del equipo: {obs}

DOCUMENTOS ACTUALES EN EL SGI:
{docs_txt}

{ISO_9001_CLAUSULAS}

Por favor proporciona:
1. **Evaluación del estado actual** de esta actividad en el SGI
2. **Qué documentos del SGI son relevantes** para esta actividad y cómo deben usarse
3. **Checklist de verificación** con 4-5 puntos concretos a revisar
4. **Alertas o riesgos** si la actividad está pendiente o vencida
5. **Próximos pasos concretos** para implementar o mejorar esta actividad

Respuesta concreta, práctica, en español, orientada a la acción."""
    return _call_gemini(prompt) or "Error al generar análisis. Verifique la API Key."


def analyze_hallazgo_redaccion(tipo: str, clausula: str, descripcion_libre: str) -> str:
    """Sugiere la redacción formal de un hallazgo de auditoría."""
    prompt = f"""Eres Auditor Líder certificado ISO 9001:2015 e ISO 39001:2015.

El auditor detectó el siguiente hallazgo durante la auditoría interna:

TIPO DE HALLAZGO: {tipo}
CLÁUSULA RELACIONADA: {clausula}
DESCRIPCIÓN LIBRE DEL AUDITOR: {descripcion_libre}

{ISO_9001_CLAUSULAS}

Proporciona:
1. **Redacción formal del hallazgo** según criterios de auditoría ISO (evidencia objetiva + criterio incumplido)
2. **Evidencia objetiva sugerida** que debería haberse recopilado
3. **Plan de acción correctiva sugerido** con responsable tipo y fecha límite estimada
4. **Cláusulas relacionadas** que podrían verse afectadas
5. **Pregunta de cierre** para verificar la eficacia de la acción

Respuesta en español, formato profesional de auditoría."""
    return _call_gemini(prompt) or "Error al generar análisis."


def extract_hallazgos_from_report(text: str, filename: str) -> list:
    """Extrae hallazgos estructurados de un informe o plan de auditoría."""
    prompt = f"""Eres un Auditor Líder certificado ISO 9001:2015 e ISO 39001:2015.

Analiza el siguiente documento de auditoría y extrae TODOS los hallazgos mencionados:
no conformidades (mayores y menores), observaciones y oportunidades de mejora.

DOCUMENTO: {filename}
CONTENIDO (primeros 6000 caracteres):
{text[:6000]}

Responde SOLO con un JSON válido con esta estructura exacta:
{{
  "hallazgos": [
    {{
      "tipo": "<No Conformidad Mayor|No Conformidad Menor|Observación|Oportunidad de Mejora>",
      "proceso": "<proceso o área auditada>",
      "clausula": "<cláusula ISO relacionada, ej: ISO 9001:2015 — 7.5.2>",
      "descripcion": "<descripción formal del hallazgo con evidencia objetiva y criterio incumplido>",
      "evidencia": "<evidencia objetiva mencionada en el documento>",
      "plan_accion": "<plan de acción si figura en el documento, caso contrario cadena vacía>",
      "responsable": "<responsable si figura, caso contrario cadena vacía>",
      "fecha_limite": "<fecha límite si figura, caso contrario cadena vacía>",
      "auditor": "<nombre del auditor si figura, caso contrario cadena vacía>"
    }}
  ]
}}

Reglas de clasificación:
- NC Mayor: incumplimiento sistémico, repetitivo o que compromete al sistema completo
- NC Menor: incumplimiento puntual o aislado de un requisito
- Observación: riesgo potencial detectado sin incumplimiento directo demostrado
- OPM: oportunidad de mejora sugerida, no es incumplimiento

Si el documento no contiene hallazgos identificables, devuelve {{"hallazgos": []}}
Extrae ABSOLUTAMENTE TODOS los hallazgos que puedas identificar en el texto."""
    r = _parse_json(_call_gemini(prompt))
    if r and "hallazgos" in r and isinstance(r["hallazgos"], list):
        return r["hallazgos"]
    return []


def generate_document_content(messages: list, base_context: str) -> dict | None:
    """Genera contenido estructurado para un documento SGI nuevo basado en el chat."""
    conv = "\n".join(
        f"{'Usuario' if m['role']=='user' else 'Asistente'}: {m['content']}"
        for m in messages
    )
    today = datetime.now().strftime("%d/%m/%Y")
    prompt = f"""Eres experto en elaboración de documentos para Sistemas de Gestión Integrados (SGI) según ISO 9001:2015 e ISO 39001:2015.

CONVERSACIÓN CON EL USUARIO:
{conv}

DOCUMENTACIÓN BASE DEL SGI (para coherencia y estilo):
{base_context[:3000]}

{ISO_9001_CLAUSULAS}

Genera un documento SGI completo, correctamente estructurado y alineado con las normas relevantes.
Responde SOLO JSON válido:
{{
  "titulo": "Título completo del documento",
  "codigo": "SGI-[TIPO]-[NRO]",
  "revision": "01",
  "fecha": "{today}",
  "organizacion": "SGI",
  "descripcion_cambio": "Emisión inicial",
  "objetivo": "Texto del objetivo (2-3 párrafos)",
  "alcance": "Texto del alcance de aplicación",
  "definiciones": [
    {{"termino": "Término", "definicion": "Definición clara"}}
  ],
  "responsabilidades": [
    {{"rol": "Cargo/Rol", "responsabilidad": "Descripción"}}
  ],
  "desarrollo": [
    {{"punto": "5.1 Subtítulo", "contenido": "Contenido detallado del punto"}},
    {{"punto": "5.2 Subtítulo", "contenido": "Contenido detallado"}}
  ],
  "referencias": ["ISO 9001:2015 — Cláusula X", "Documento relacionado"]
}}"""
    return _parse_json(_call_gemini(prompt))


# ─── RAG QUERY ────────────────────────────────────────────────────────────────
def rag_query(question: str) -> str:
    col = get_collection()
    total = col.count()
    if total == 0:
        return "No hay documentos indexados. Suba archivos primero."
    try:
        results = col.query(query_texts=[question], n_results=min(RAG_TOP_K, total))
        docs, metas = results["documents"][0], results["metadatas"][0]
    except Exception as e:
        return f"Error en búsqueda vectorial: {e}"
    if not docs: return "No se encontraron fragmentos relevantes."
    context = "\n\n---\n\n".join(f"[{m.get('source','?')}]\n{d}" for d, m in zip(docs, metas))
    if not os.environ.get("GEMINI_API_KEY"):
        return "**Fragmentos encontrados (configure GEMINI_API_KEY para respuesta interpretada):**\n\n" + context
    prompt = f"""Eres Asistente de Auditoría SGI experto en ISO 9001:2015 e ISO 39001:2015.
Responde basándote EXCLUSIVAMENTE en estos fragmentos. Si la info no está, indícalo.

FRAGMENTOS:
{context}

PREGUNTA: {question}"""
    answer = _call_gemini(prompt)
    if answer:
        sources = ", ".join({m.get("source","?") for m in metas})
        return f"{answer}\n\n*📁 Fuentes: {sources}*"
    return "Error generando respuesta."


# ─── CHROMADB INDEXING ────────────────────────────────────────────────────────
def index_document(file_bytes: bytes, filename: str, text: str) -> None:
    fhash = hashlib.md5(file_bytes).hexdigest()[:8]
    chunks = chunk_text(text)
    if not chunks: return
    col = get_collection()
    ids  = [f"{fhash}_{i}" for i in range(len(chunks))]
    meta = [{"source":filename,"chunk_idx":i,"hash":fhash} for i in range(len(chunks))]
    try: col.upsert(documents=chunks, ids=ids, metadatas=meta)
    except Exception as e: st.warning(f"ChromaDB: {e}")


# ─── GESTIÓN DE INCONGRUENCIAS ────────────────────────────────────────────────
def migrate_incongruencias(analisis_cache: dict) -> dict:
    """Migra incongruencias del cache a items individuales rastreables."""
    items = {}
    for doc_name, analysis in analisis_cache.items():
        if isinstance(analysis, dict):
            for texto in analysis.get("incongruencias", []):
                if not texto or any(p in texto for p in ["⚠️","📌","Error","Configure"]):
                    continue
                inc_id = hashlib.md5(f"{doc_name}:{texto}".encode()).hexdigest()[:10]
                if inc_id not in items:
                    items[inc_id] = {
                        "id": inc_id, "documento": doc_name, "texto": texto,
                        "estado": "abierta",
                        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "chat": []
                    }
    return items

def load_incongruencias(analisis_cache: dict) -> dict:
    data = load_json(INCONGRUENCIAS_PATH, None)
    if data is None:
        items = migrate_incongruencias(analisis_cache)
        save_json(INCONGRUENCIAS_PATH, items)
        return items
    return data

def save_incongruencias(items: dict):
    save_json(INCONGRUENCIAS_PATH, items)

def add_incongruencia_from_analysis(doc_name: str, inc_texto: str, items: dict) -> dict:
    """Agrega una incongruencia nueva al tracker."""
    inc_id = hashlib.md5(f"{doc_name}:{inc_texto}".encode()).hexdigest()[:10]
    if inc_id not in items:
        items[inc_id] = {
            "id": inc_id, "documento": doc_name, "texto": inc_texto,
            "estado": "abierta",
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "chat": []
        }
    return items

def _incongruencia_ya_existe(texto: str, existing: dict, current_doc: str) -> bool:
    """Retorna True si ya existe una incongruencia muy similar en CUALQUIER otro documento.
    Evita registrar duplicados cross-documento (≥70 % palabras en común)."""
    words_new = set(texto.lower().split())
    if len(words_new) < 5:
        return False
    for inc in existing.values():
        if inc.get("documento") == current_doc:
            continue
        if inc.get("estado") == "eliminada":
            continue
        words_other = set(inc.get("texto", "").lower().split())
        if not words_other:
            continue
        overlap = len(words_new & words_other) / max(len(words_new), len(words_other))
        if overlap >= 0.70:
            return True
    return False


# ─── GENERACIÓN DE DOCUMENTO DOCX ────────────────────────────────────────────
def build_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(3.5); sec.bottom_margin = Cm(3.0)

    BLUE = RGBColor(13, 71, 161)
    DARK = RGBColor(30, 41, 59)

    def border_table(tbl):
        tbl_pr = tbl._tbl.get_or_add_tblPr()
        borders = OxmlElement("w:tblBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
            el.set(qn("w:color"),"1565C0"); borders.append(el)
        tbl_pr.append(borders)

    def cell_center(cell, txt, bold=False, size=8, color=DARK):
        for p in cell.paragraphs: p.clear()
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt); run.font.size = Pt(size)
        run.font.bold = bold; run.font.color.rgb = color

    def sec_title(doc, txt):
        p = doc.add_paragraph()
        run = p.add_run(txt); run.font.bold = True
        run.font.size = Pt(11); run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)

    def body_text(doc, txt):
        p = doc.add_paragraph(txt if txt else "")
        for r in p.runs: r.font.size = Pt(10); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(6)

    today = data.get("fecha", datetime.now().strftime("%d/%m/%Y"))

    # ── ENCABEZADO DE PÁGINA ──────────────────────────────────────────────────
    hdr = sec.header
    for p in hdr.paragraphs: p.clear()
    ht = hdr.add_table(rows=2, cols=4, width=Inches(6.5))
    border_table(ht)
    labels = ["ORGANIZACIÓN", "TÍTULO DEL DOCUMENTO", "CÓDIGO", "REV / FECHA"]
    vals   = [data.get("organizacion","SGI"), data.get("titulo",""), data.get("codigo","SGI-XXX-001"),
              f"Rev {data.get('revision','01')}  |  {today}"]
    for i,lbl in enumerate(labels): cell_center(ht.rows[0].cells[i], lbl, bold=True, color=BLUE)
    for i,val in enumerate(vals):   cell_center(ht.rows[1].cells[i], val, size=9)
    hdr.add_paragraph()   # espaciado

    # ── TÍTULO PRINCIPAL ──────────────────────────────────────────────────────
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(data.get("titulo","DOCUMENTO SGI").upper())
    tr.font.size = Pt(14); tr.font.bold = True; tr.font.color.rgb = BLUE
    tp.paragraph_format.space_after = Pt(16)

    # ── SECCIONES ─────────────────────────────────────────────────────────────
    sec_title(doc, "1.  OBJETIVO")
    body_text(doc, data.get("objetivo",""))

    sec_title(doc, "2.  ALCANCE")
    body_text(doc, data.get("alcance",""))

    defs = data.get("definiciones", [])
    if defs:
        sec_title(doc, "3.  DEFINICIONES Y ABREVIATURAS")
        for d in defs:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{d.get('termino','')}: "); r1.font.bold = True; r1.font.size = Pt(10)
            r2 = p.add_run(d.get("definicion","")); r2.font.size = Pt(10); r2.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(4)

    resps = data.get("responsabilidades", [])
    if resps:
        sec_title(doc, "4.  RESPONSABILIDADES")
        for r in resps:
            p = doc.add_paragraph(style="List Bullet")
            r1 = p.add_run(f"{r.get('rol','')}: "); r1.font.bold = True; r1.font.size = Pt(10)
            r2 = p.add_run(r.get("responsabilidad","")); r2.font.size = Pt(10); r2.font.color.rgb = DARK

    desarr = data.get("desarrollo", [])
    if desarr:
        sec_title(doc, "5.  DESARROLLO")
        for item in desarr:
            p = doc.add_paragraph()
            sh = p.add_run(item.get("punto",""))
            sh.font.bold = True; sh.font.size = Pt(10); sh.font.color.rgb = DARK
            body_text(doc, item.get("contenido",""))

    refs = data.get("referencias", [])
    if refs:
        sec_title(doc, "6.  REFERENCIAS Y DOCUMENTOS RELACIONADOS")
        for r in refs:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(r); run.font.size = Pt(10); run.font.color.rgb = DARK

    # Control de cambios
    sec_title(doc, "7.  CONTROL DE CAMBIOS")
    ct = doc.add_table(rows=2, cols=4); border_table(ct)
    for i, h in enumerate(["Revisión","Fecha","Descripción del Cambio","Responsable"]):
        cell_center(ct.rows[0].cells[i], h, bold=True, color=BLUE)
    vals_ct = [data.get("revision","01"), today, data.get("descripcion_cambio","Emisión inicial"), "_______________"]
    for i, v in enumerate(vals_ct): cell_center(ct.rows[1].cells[i], v, size=9)

    # ── PIE DE PÁGINA ─────────────────────────────────────────────────────────
    ftr = sec.footer
    for p in ftr.paragraphs: p.clear()
    ft = ftr.add_table(rows=4, cols=3, width=Inches(6.5)); border_table(ft)
    for i, role in enumerate(["ELABORADO POR","REVISADO POR","APROBADO POR"]):
        cell_center(ft.rows[0].cells[i], role, bold=True, color=BLUE, size=8)
    for row_i, lbl in enumerate(["Nombre:  ___________________________",
                                  "Firma:     ___________________________",
                                  "Fecha:    ___________________________"], 1):
        for col_i in range(3):
            cell_center(ft.rows[row_i].cells[col_i], lbl, size=8)

    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()


# ─── UI HELPERS ───────────────────────────────────────────────────────────────
BADGE_MAP = {"Procedimiento":"badge-proc","Instructivo":"badge-inst","Registro":"badge-reg",
             "Manual":"badge-man","Política":"badge-pol","Plan":"badge-pol",
             "Informe":"badge-inf","Formulario":"badge-reg"}

def tipo_badge(tipo: str) -> str:
    cls = BADGE_MAP.get(tipo, "badge-other")
    return f'<span class="doc-badge {cls}">{tipo}</span>'

def status_pill(estado: str) -> str:
    icons = {"abierta":"🔴","resuelta":"✅","eliminada":"⚫"}
    ico = icons.get(estado, "•")
    return f'<span class="status-pill st-{estado}">{ico} {estado.capitalize()}</span>'

def verdict_html(veredicto: str) -> str:
    v = veredicto.upper()
    if "APROBADO" in v: return f'<span class="verdict verdict-ok">✅ {veredicto}</span>'
    if "OBSERV" in v:   return f'<span class="verdict verdict-obs">⚠️ {veredicto}</span>'
    return f'<span class="verdict verdict-nok">❌ {veredicto}</span>'

def step_html(n, text: str) -> str:
    return (f'<div class="step-item"><div class="step-num">{n}</div>'
            f'<div class="step-txt">{text}</div></div>')


# ─── TAB: LISTA MAESTRA / MATRIZ NORMATIVA ───────────────────────────────────
def tab_lista_maestra(lista_maestra: list):
    activos   = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
    obsoletos = [d for d in lista_maestra if d.get("estado","activo")=="obsoleto"]

    st.markdown('<p class="sec-title">📋 Lista Maestra y Matriz Normativa</p>', unsafe_allow_html=True)

    if not lista_maestra:
        st.markdown('<div class="card-info" style="color:#1e293b!important">Aún no hay documentos '
                    'en el SGI. Suba documentos desde la pestaña <b>📂 Documentos</b>.</div>',
                    unsafe_allow_html=True)
        return

    # KPIs rápidos
    k1,k2,k3,k4 = st.columns(4)
    tipos = {}
    for d in activos: tipos[d.get("tipo","Otro")] = tipos.get(d.get("tipo","Otro"),0)+1
    with k1: st.metric("Documentos activos", len(activos))
    with k2: st.metric("Tipos diferentes", len(tipos))
    with k3: st.metric("Obsoletos", len(obsoletos))
    with k4: st.metric("Total en sistema", len(lista_maestra))

    st.markdown("---")

    # Subtabs
    st_lm, st_mat = st.tabs(["📄  Lista Maestra Completa","🗺️  Matriz Doc ↔ Cláusula"])

    # ── LISTA MAESTRA ─────────────────────────────────────────────────────────
    with st_lm:
        lf1, lf2, lf3 = st.columns([2, 2, 2])
        with lf1:
            filtro_tipo = st.selectbox("Tipo de documento",
                                       ["Todos"] + sorted(set(d.get("tipo","Otro") for d in lista_maestra)),
                                       key="lm_tipo")
        with lf2:
            filtro_est = st.selectbox("Estado", ["activo","obsoleto","Todos"], key="lm_est")
        with lf3:
            filtro_norma = st.selectbox("Cobertura normativa",
                                        ["Todas","ISO 9001:2015","ISO 39001:2015","Ambas normas"],
                                        key="lm_norma")

        filas = lista_maestra
        if filtro_tipo != "Todos":
            filas = [d for d in filas if d.get("tipo") == filtro_tipo]
        if filtro_est != "Todos":
            filas = [d for d in filas if d.get("estado","activo") == filtro_est]
        if filtro_norma == "ISO 9001:2015":
            filas = [d for d in filas if d.get("iso9001","") and
                     "no aplica" not in d.get("iso9001","").lower()]
        elif filtro_norma == "ISO 39001:2015":
            filas = [d for d in filas if d.get("iso39001","") and
                     "no aplica" not in d.get("iso39001","").lower()]
        elif filtro_norma == "Ambas normas":
            filas = [d for d in filas if
                     d.get("iso9001","") and "no aplica" not in d.get("iso9001","").lower() and
                     d.get("iso39001","") and "no aplica" not in d.get("iso39001","").lower()]

        if not filas:
            st.info("No hay documentos con los filtros seleccionados.")

        if filas:
            # Tabla HTML
            rows_html = ""
            for d in sorted(filas, key=lambda x: x.get("nombre","")):
                est = d.get("estado","activo")
                est_html = (
                    '<span class="status-pill" style="background:#fee2e2;color:#b91c1c">⚫ Obsoleto</span>'
                    if est == "obsoleto"
                    else '<span class="status-pill" style="background:#d1fae5;color:#065f46">✅ Activo</span>'
                )
                iso9_txt  = d.get("iso9001","—")[:80]
                iso39_txt = d.get("iso39001","—")[:80]
                rows_html += (
                    f"<tr>"
                    f"<td><b>{d.get('nombre','')}</b><br>"
                    f"<span style='font-size:.69rem;color:#9ca3af'>{d.get('fecha_ingreso','')}</span></td>"
                    f"<td>{tipo_badge(d.get('tipo','Otro'))}</td>"
                    f"<td style='font-size:.78rem'>{iso9_txt}</td>"
                    f"<td style='font-size:.78rem'>{iso39_txt}</td>"
                    f"<td>{est_html}</td>"
                    f"</tr>"
                )
            st.markdown(
                f'<table class="matriz-table"><thead><tr>'
                f'<th>Documento</th><th>Tipo</th>'
                f'<th>ISO 9001:2015</th><th>ISO 39001:2015</th><th>Estado</th>'
                f'</tr></thead><tbody>{rows_html}</tbody></table>',
                unsafe_allow_html=True,
            )
            st.markdown("<br>", unsafe_allow_html=True)
            # Descarga CSV
            import io, csv
            buf = io.StringIO()
            w = csv.writer(buf)
            w.writerow(["Nombre","Tipo","ISO 9001","ISO 39001","Estado","Fecha Ingreso"])
            for d in filas:
                w.writerow([d.get("nombre",""), d.get("tipo",""), d.get("iso9001",""),
                            d.get("iso39001",""), d.get("estado",""), d.get("fecha_ingreso","")])
            st.download_button("⬇️ Descargar Lista Maestra (CSV)", buf.getvalue().encode(),
                               "lista_maestra_sgi.csv", "text/csv", key="dl_lm_csv")

        # ── Re-analizar documentos con Error ──────────────────────────────────
        docs_con_error = [
            d for d in lista_maestra
            if d.get("iso9001","") in ("Error","Sin clasificar","") or
               d.get("iso39001","") in ("Error","Sin clasificar","") or
               d.get("tipo","") in ("Sin clasificar","")
        ]
        if docs_con_error:
            st.markdown("---")
            with st.expander(f"🔄  Re-analizar {len(docs_con_error)} documento(s) con Error / Sin clasificar",
                             expanded=True):
                st.markdown(
                    '<div class="card-warn" style="color:#1e293b!important;font-size:.83rem">'
                    'Estos documentos quedaron sin clasificar porque la API de Gemini no estaba disponible '
                    'cuando se subieron. Hacé clic en <b>Re-analizar</b> para procesarlos ahora.</div>',
                    unsafe_allow_html=True,
                )
                # Botón "Re-analizar TODOS" con pausa automática
                import time as _time
                ra_col1, ra_col2 = st.columns([2, 3])
                with ra_col1:
                    if st.button("🔄 Re-analizar TODOS (automático)",
                                 key="reanalyze_all_btn", use_container_width=True,
                                 type="primary"):
                        st.session_state["do_reanalyze_all"] = True

                with ra_col2:
                    st.markdown(
                        '<div class="card-info" style="color:#1e293b!important;'
                        'font-size:.79rem;padding:8px 12px">'
                        '💡 Procesa uno por uno con pausa de 5 s para no agotar cuota. '
                        'O configurá <b>GROQ_API_KEY</b> para análisis sin límite.</div>',
                        unsafe_allow_html=True,
                    )

                st.markdown("<br>", unsafe_allow_html=True)

                for doc in docs_con_error:
                    ec1, ec2 = st.columns([5, 1])
                    with ec1:
                        st.markdown(
                            f'<div style="font-size:.87rem;color:#1e293b;padding:4px 0">'
                            f'📄 <b>{doc["nombre"]}</b>'
                            f'<span style="font-size:.75rem;color:#9ca3af;margin-left:8px">'
                            f'ISO 9001: {doc.get("iso9001","—")[:40]} · '
                            f'ISO 39001: {doc.get("iso39001","—")[:40]}</span></div>',
                            unsafe_allow_html=True,
                        )
                    with ec2:
                        if st.button("🔄", key=f"reanalyze_{doc['hash']}",
                                     use_container_width=True,
                                     help=f"Re-analizar solo este documento"):
                            st.session_state[f"do_reanalyze_{doc['hash']}"] = True

                # ── Función auxiliar para re-analizar un documento ──────────
                def _do_reanalyze_doc(doc, lista_maestra, analisis_cache_local):
                    doc_text = (analisis_cache_local.get(doc["nombre"], {}).get("_texto", "")
                                or get_document_text_from_chroma(doc["nombre"]))
                    if not doc_text:
                        st.error(f"❌ Sin texto para **{doc['nombre']}**. "
                                 "Eliminalo y volvé a subirlo.")
                        return False
                    with st.spinner(f"Analizando **{doc['nombre'][:50]}**..."):
                        new_cls  = classify_document(doc_text, doc["nombre"])
                        new_deep = analyze_document_deep(doc_text, doc["nombre"])
                    ok = new_cls.get("iso9001","") not in ("Error","Configure GEMINI_API_KEY","")
                    if ok:
                        for i, d in enumerate(lista_maestra):
                            if d.get("hash") == doc["hash"]:
                                lista_maestra[i]["tipo"]     = new_cls.get("tipo", d["tipo"])
                                lista_maestra[i]["iso9001"]  = new_cls.get("iso9001", d["iso9001"])
                                lista_maestra[i]["iso39001"] = new_cls.get("iso39001", d["iso39001"])
                                break
                        save_json(LISTA_MAESTRA_PATH, lista_maestra)
                        analisis_cache_local[doc["nombre"]] = {
                            **new_deep, "_texto": doc_text[:TEXT_PREVIEW_LEN]}
                        save_json(ANALISIS_PATH, analisis_cache_local)
                        st.success(f"✅ **{doc['nombre'][:50]}** re-analizado.")
                    else:
                        st.warning(f"⏳ Sin respuesta para **{doc['nombre'][:50]}**.")
                    return ok

                # Ejecutar re-análisis de TODOS
                analisis_cache_local = load_json(ANALISIS_PATH, {})
                if st.session_state.get("do_reanalyze_all"):
                    st.session_state.pop("do_reanalyze_all", None)
                    total = len(docs_con_error)
                    prog_bar = st.progress(0, text=f"Re-analizando 0/{total}...")
                    for idx, doc in enumerate(docs_con_error):
                        prog_bar.progress((idx) / total,
                                          text=f"Re-analizando {idx+1}/{total}: {doc['nombre'][:40]}...")
                        _do_reanalyze_doc(doc, lista_maestra, analisis_cache_local)
                        if idx < total - 1:
                            _time.sleep(5)  # 5 s de pausa entre docs para no agotar cuota
                    prog_bar.progress(1.0, text="✅ Re-análisis completado.")
                    st.rerun()

                # Ejecutar re-análisis individual
                for doc in docs_con_error:
                    if st.session_state.get(f"do_reanalyze_{doc['hash']}"):
                        st.session_state.pop(f"do_reanalyze_{doc['hash']}", None)
                        _do_reanalyze_doc(doc, lista_maestra, analisis_cache_local)
                        _time.sleep(2)
                        st.rerun()

        # ── Eliminar documento ─────────────────────────────────────────────────
        st.markdown("---")
        with st.expander("🗑️  Eliminar documento de la Lista Maestra", expanded=False):
            st.markdown(
                '<div class="card-warn" style="color:#1e293b!important;font-size:.83rem">'
                '⚠️ Al eliminar un documento se borra de la Lista Maestra y se eliminan '
                'sus incongruencias asociadas. Esta acción no se puede deshacer.</div>',
                unsafe_allow_html=True,
            )
            doc_nombres_todos = [d["nombre"] for d in sorted(lista_maestra, key=lambda x: x.get("nombre",""))]
            del_c1, del_c2 = st.columns([5, 1])
            with del_c1:
                doc_to_delete = st.selectbox(
                    "Documento a eliminar:",
                    ["— Seleccioná un documento —"] + doc_nombres_todos,
                    key="lm_del_selectbox",
                    label_visibility="collapsed",
                )
            with del_c2:
                del_btn = st.button("🗑️ Eliminar", key="lm_del_btn",
                                    disabled=(doc_to_delete == "— Seleccioná un documento —"),
                                    use_container_width=True)

            if del_btn and doc_to_delete != "— Seleccioná un documento —":
                st.session_state["lm_confirm_del"] = doc_to_delete

            if st.session_state.get("lm_confirm_del"):
                cname = st.session_state["lm_confirm_del"]
                st.error(f"❗ ¿Eliminar definitivamente **{cname}**?")
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button("✅ Sí, eliminar", key="lm_del_confirm_yes", type="primary",
                                 use_container_width=True):
                        # Eliminar de lista maestra
                        nueva_lista = [d for d in lista_maestra if d["nombre"] != cname]
                        save_json(LISTA_MAESTRA_PATH, nueva_lista)
                        # Eliminar incongruencias del doc
                        incs = load_json(INCONGRUENCIAS_PATH, {})
                        incs = {k: v for k, v in incs.items() if v.get("documento") != cname}
                        save_json(INCONGRUENCIAS_PATH, incs)
                        # Eliminar del analisis cache
                        ac = load_json(ANALISIS_PATH, {})
                        ac.pop(cname, None)
                        save_json(ANALISIS_PATH, ac)
                        st.session_state.pop("lm_confirm_del", None)
                        st.success(f"✅ **{cname}** eliminado correctamente.")
                        st.rerun()
                with cc2:
                    if st.button("❌ Cancelar", key="lm_del_cancel_btn", use_container_width=True):
                        st.session_state.pop("lm_confirm_del", None)
                        st.rerun()

    # ── MATRIZ DOC ↔ CLÁUSULA ─────────────────────────────────────────────────
    with st_mat:
        norma_sel = st.radio("Norma:", ["ISO 9001:2015", "ISO 39001:2015", "🗺️ Ambas normas"],
                             horizontal=True, key="mat_norma")

        docs_act = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
        if not docs_act:
            st.info("No hay documentos activos para mostrar la matriz.")
            return

        # ── Vista AMBAS NORMAS ───────────────────────────────────────────────
        if "Ambas" in norma_sel:
            st.markdown(
                '<div class="card-info" style="color:#1e293b!important;font-size:.84rem">'
                'Vista combinada: cobertura de <b>ISO 9001:2015</b> (azul) e '
                '<b>ISO 39001:2015</b> (violeta) en una sola tabla.</div>',
                unsafe_allow_html=True,
            )

            def _cobertura_tags(clausula, field, css_class, docs):
                found = [d["nombre"] for d in docs if clausula in d.get(field,"")]
                if not found:
                    return '<span style="color:#94a3b8;font-size:.73rem">⚠️ Sin cobertura</span>', False
                tags = "".join(
                    f'<span class="clausula-tag {css_class}">'
                    f'{n[:38]}{"…" if len(n)>38 else ""}</span>'
                    for n in found
                )
                return tags, True

            rows_mat = ""
            clausulas_all = sorted(
                set(list(ISO_9001_DICT.keys()) + list(ISO_39001_DICT.keys())),
                key=lambda x: [int(p) for p in x.split(".")]
            )
            for clausula in clausulas_all:
                nombre_9001  = ISO_9001_DICT.get(clausula, "")
                nombre_39001 = ISO_39001_DICT.get(clausula, "")
                tags_9001, ok_9001   = _cobertura_tags(clausula, "iso9001",  "",     docs_act)
                tags_39001, ok_39001 = _cobertura_tags(clausula, "iso39001", "clausula-tag-39", docs_act)

                norma_label = ""
                if nombre_9001  and nombre_39001:
                    norma_label = '<span style="font-size:.68rem;color:#1d4ed8">9001</span> / <span style="font-size:.68rem;color:#7c3aed">39001</span>'
                elif nombre_9001:
                    norma_label = '<span style="font-size:.68rem;color:#1d4ed8">solo 9001</span>'
                else:
                    norma_label = '<span style="font-size:.68rem;color:#7c3aed">solo 39001</span>'

                nombre_clausula = nombre_9001 or nombre_39001
                fondo = "" if (ok_9001 or ok_39001) else 'style="background:#fff8f8"'

                rows_mat += (
                    f"<tr {fondo}>"
                    f"<td style='font-weight:700;color:#1d4ed8;white-space:nowrap'>"
                    f"{clausula}<br>{norma_label}</td>"
                    f"<td style='color:#374151;font-size:.82rem'>{nombre_clausula}</td>"
                    f"<td style='font-size:.78rem'><b style='color:#1d4ed8'>9001:</b><br>{tags_9001}</td>"
                    f"<td style='font-size:.78rem'><b style='color:#7c3aed'>39001:</b><br>{tags_39001}</td>"
                    f"</tr>"
                )

            st.markdown(
                f'<table class="matriz-table"><thead><tr>'
                f'<th style="width:90px">Cláusula</th>'
                f'<th style="width:210px">Nombre</th>'
                f'<th>Cobertura ISO 9001:2015</th>'
                f'<th>Cobertura ISO 39001:2015</th>'
                f'</tr></thead><tbody>{rows_mat}</tbody></table>',
                unsafe_allow_html=True,
            )

            # KPIs de cobertura de ambas normas
            def _pct(iso_dict, field):
                total = len(iso_dict)
                cub   = sum(1 for c in iso_dict if any(c in d.get(field,"") for d in docs_act))
                return cub, total, int(cub/total*100) if total else 0

            c9001, t9001, p9001 = _pct(ISO_9001_DICT, "iso9001")
            c3900, t3900, p3900 = _pct(ISO_39001_DICT, "iso39001")
            col_a, col_b = st.columns(2)
            for col, pct, cub, tot, label, color_scheme in [
                (col_a, p9001, c9001, t9001, "ISO 9001:2015", "#1d4ed8"),
                (col_b, p3900, c3900, t3900, "ISO 39001:2015", "#7c3aed"),
            ]:
                color = "#22c55e" if pct>=80 else "#f59e0b" if pct>=50 else "#ef4444"
                with col:
                    st.markdown(
                        f'<div class="card" style="margin-top:14px;text-align:center">'
                        f'<b style="color:{color_scheme}">{label}</b><br>'
                        f'<span style="font-size:1.6rem;font-weight:800;color:{color}">{pct}%</span>'
                        f'<br><span style="font-size:.8rem;color:#6b7280">'
                        f'{cub}/{tot} cláusulas cubiertas</span></div>',
                        unsafe_allow_html=True,
                    )

        # ── Vista NORMA INDIVIDUAL ───────────────────────────────────────────
        else:
            iso_dict = ISO_9001_DICT if "9001" in norma_sel else ISO_39001_DICT
            field    = "iso9001" if "9001" in norma_sel else "iso39001"
            css_cls  = "" if "9001" in norma_sel else "clausula-tag-39"

            st.markdown(
                f'<div class="card-info" style="color:#1e293b!important;font-size:.84rem">'
                f'La matriz muestra qué documentos del SGI cubren cada cláusula de <b>{norma_sel}</b>. '
                f'Los documentos clasificados por IA aparecen automáticamente.</div>',
                unsafe_allow_html=True,
            )

            rows_mat = ""
            for clausula, nombre_clausula in iso_dict.items():
                docs_que_cubren = [d["nombre"] for d in docs_act if clausula in d.get(field,"")]
                if docs_que_cubren:
                    tags = "".join(
                        f'<span class="clausula-tag {css_cls}">'
                        f'{d[:45]}{"…" if len(d)>45 else ""}</span>'
                        for d in docs_que_cubren
                    )
                    cobertura = tags; fondo = ""
                else:
                    cobertura = '<span style="color:#94a3b8;font-size:.75rem">⚠️ Sin cobertura documental</span>'
                    fondo = 'style="background:#fff8f8"'

                rows_mat += (
                    f"<tr {fondo}>"
                    f"<td style='font-weight:700;color:#1d4ed8;white-space:nowrap'>{clausula}</td>"
                    f"<td style='color:#374151'>{nombre_clausula}</td>"
                    f"<td>{cobertura}</td>"
                    f"</tr>"
                )

            st.markdown(
                f'<table class="matriz-table"><thead><tr>'
                f'<th style="width:80px">Cláusula</th>'
                f'<th style="width:260px">Nombre</th>'
                f'<th>Documentos del SGI que la cubren</th>'
                f'</tr></thead><tbody>{rows_mat}</tbody></table>',
                unsafe_allow_html=True,
            )

            total_clausulas = len(iso_dict)
            con_cobertura   = sum(1 for c in iso_dict if any(c in d.get(field,"") for d in docs_act))
            pct   = int(con_cobertura / total_clausulas * 100)
            color = "#22c55e" if pct>=80 else "#f59e0b" if pct>=50 else "#ef4444"
            st.markdown(
                f'<div class="card" style="margin-top:14px;text-align:center">'
                f'<b>Cobertura normativa {norma_sel}:</b>&nbsp;&nbsp;'
                f'<span style="font-size:1.4rem;font-weight:800;color:{color}">{pct}%</span>'
                f'&nbsp; ({con_cobertura}/{total_clausulas} cláusulas cubiertas por documentos)</div>',
                unsafe_allow_html=True,
            )


# ─── TAB: AGENDA SGI ─────────────────────────────────────────────────────────
def tab_agenda_sgi(lista_maestra: list):
    agenda = load_agenda()
    activos = [d for d in lista_maestra if d.get("estado","activo")=="activo"]

    st.markdown('<p class="sec-title">📅 Agenda de Actividades Obligatorias del SGI</p>',
                unsafe_allow_html=True)
    st.markdown(
        '<div class="card-info" style="color:#1e293b!important;font-size:.84rem">'
        'Registro de las actividades que el SGI debe realizar periódicamente según ISO 9001:2015 '
        'e ISO 39001:2015. Marcá cada actividad cuando se realice y solicitá análisis de IA '
        'para obtener orientación específica.</div>',
        unsafe_allow_html=True,
    )

    # Filtros
    fc1, fc2 = st.columns([3,3])
    with fc1:
        filt_freq = st.selectbox("Frecuencia", ["Todas","Mensual","Trimestral","Semestral","Anual"],
                                 key="ag_freq")
    with fc2:
        filt_est = st.selectbox("Estado", ["Todas","pendiente","realizado"], key="ag_est")

    actividades = agenda
    if filt_freq != "Todas": actividades = [a for a in actividades if a["frecuencia"]==filt_freq]
    if filt_est  != "Todas": actividades = [a for a in actividades if a["estado"]==filt_est]

    # KPIs
    pendientes  = sum(1 for a in agenda if a["estado"]=="pendiente")
    realizados  = sum(1 for a in agenda if a["estado"]=="realizado")
    ag1,ag2,ag3 = st.columns(3)
    with ag1: st.metric("Pendientes", pendientes)
    with ag2: st.metric("Realizados", realizados)
    with ag3: st.metric("Total actividades", len(agenda))

    st.markdown("---")

    FREQ_COLORS = {"Mensual":"freq-pill-mes","Trimestral":"freq-pill-trim",
                   "Semestral":"freq-pill","Anual":"freq-pill-anual"}

    for act in actividades:
        est  = act.get("estado","pendiente")
        done = est == "realizado"
        card_cls = "agenda-card agenda-card-done" if done else "agenda-card"

        st.markdown(
            f'<div class="{card_cls}">'
            f'<div style="display:flex;justify-content:space-between;align-items:flex-start;gap:8px">'
            f'<div style="font-weight:700;font-size:.93rem;color:#1a237e">{act["nombre"]}</div>'
            f'<span class="freq-pill {FREQ_COLORS.get(act["frecuencia"],"freq-pill")}">'
            f'{act["frecuencia"]}</span></div>'
            f'<div style="font-size:.81rem;color:#6b7280;margin:5px 0 4px">'
            f'📌 9001: <b>{act["clausula_9001"]}</b> &nbsp;·&nbsp; '
            f'39001: <b>{act["clausula_39001"]}</b> &nbsp;·&nbsp; '
            f'👤 {act["responsable"]}</div>'
            f'<div style="font-size:.84rem;color:#374151;line-height:1.5">{act["descripcion"]}</div>'
            + (f'<div style="font-size:.78rem;color:#15803d;margin-top:6px">✅ Realizado: '
               f'{act.get("fecha_realizacion","")} &nbsp;·&nbsp; {act.get("observaciones","")}</div>'
               if done else "")
            + '</div>',
            unsafe_allow_html=True,
        )

        ba1, ba2, ba3 = st.columns([2,2,3])
        with ba1:
            if not done:
                if st.button("✅ Marcar realizado", key=f"ag_done_{act['id']}", use_container_width=True):
                    st.session_state[f"ag_form_{act['id']}"] = True
            else:
                if st.button("🔄 Marcar pendiente", key=f"ag_pend_{act['id']}", use_container_width=True):
                    act["estado"]="pendiente"; act["fecha_realizacion"]=""
                    idx = next(i for i,a in enumerate(agenda) if a["id"]==act["id"])
                    agenda[idx] = act; save_json(AGENDA_PATH, agenda); st.rerun()
        with ba2:
            if st.button("🤖 Análisis IA", key=f"ag_ai_{act['id']}", use_container_width=True):
                st.session_state[f"ag_ai_show_{act['id']}"] = True
        with ba3:
            fecha_lim = st.text_input("Fecha límite (YYYY-MM-DD)",
                                       value=act.get("fecha_limite",""),
                                       key=f"ag_fl_{act['id']}",
                                       label_visibility="collapsed",
                                       placeholder="📆 Fecha límite (YYYY-MM-DD)")
            if fecha_lim != act.get("fecha_limite",""):
                act["fecha_limite"] = fecha_lim
                idx = next(i for i,a in enumerate(agenda) if a["id"]==act["id"])
                agenda[idx] = act; save_json(AGENDA_PATH, agenda)

        # Formulario "marcar realizado"
        if st.session_state.get(f"ag_form_{act['id']}"):
            with st.form(key=f"form_realizado_{act['id']}"):
                obs_inp = st.text_area("Observaciones (opcional)", key=f"ag_obs_{act['id']}")
                fecha_r = st.text_input("Fecha de realización", value=datetime.now().strftime("%Y-%m-%d"))
                if st.form_submit_button("Confirmar"):
                    act["estado"]="realizado"; act["fecha_realizacion"]=fecha_r
                    act["observaciones"]=obs_inp
                    idx = next(i for i,a in enumerate(agenda) if a["id"]==act["id"])
                    agenda[idx] = act; save_json(AGENDA_PATH, agenda)
                    st.session_state.pop(f"ag_form_{act['id']}", None); st.rerun()

        # Panel de análisis IA
        if st.session_state.get(f"ag_ai_show_{act['id']}"):
            with st.spinner(f"Analizando actividad «{act['nombre']}»..."):
                feedback = analyze_sgi_activity(act, activos)
            act["ai_feedback"] = feedback
            idx = next(i for i,a in enumerate(agenda) if a["id"]==act["id"])
            agenda[idx] = act; save_json(AGENDA_PATH, agenda)
            st.session_state.pop(f"ag_ai_show_{act['id']}", None)

        if act.get("ai_feedback"):
            with st.expander("🤖 Análisis IA — clic para ver/ocultar", expanded=False):
                st.markdown(act["ai_feedback"])

        st.markdown("", unsafe_allow_html=True)


# ─── TAB: AUDITORÍA INTERNA ───────────────────────────────────────────────────
TIPOS_HALLAZGO = ["No Conformidad Mayor","No Conformidad Menor","Observación","Oportunidad de Mejora"]
ESTADOS_HALLAZGO = ["abierto","en_proceso","cerrado","verificado"]
HTIPO_CSS = {"No Conformidad Mayor":"htipo-ncm","No Conformidad Menor":"htipo-ncn",
             "Observación":"htipo-obs","Oportunidad de Mejora":"htipo-opm"}
HCARD_CSS = {"No Conformidad Mayor":"hallazgo-nc-mayor","No Conformidad Menor":"hallazgo-nc-menor",
             "Observación":"hallazgo-obs","Oportunidad de Mejora":"hallazgo-opm"}

def hallazgo_tipo_html(tipo: str) -> str:
    cls = HTIPO_CSS.get(tipo,"htipo-obs")
    return f'<span class="htipo-pill {cls}">{tipo}</span>'

def hallazgo_estado_html(estado: str) -> str:
    cls = f"hestado-pill he-{estado}"
    icons = {"abierto":"🔴","en_proceso":"🟡","cerrado":"✅","verificado":"🔵"}
    return f'<span class="{cls}">{icons.get(estado,"•")} {estado.replace("_"," ").capitalize()}</span>'

def tab_auditoria(lista_maestra: list, analisis_cache: dict):
    hallazgos = load_hallazgos()
    activos   = [d for d in lista_maestra if d.get("estado","activo")=="activo"]

    st.markdown('<p class="sec-title">🔍 Auditoría Interna — Registro de Hallazgos</p>',
                unsafe_allow_html=True)

    ht1, ht2, ht3 = st.tabs(["📋  Registro de Hallazgos","➕  Nuevo Hallazgo","📤  Importar desde Informe"])

    # ── LISTADO ───────────────────────────────────────────────────────────────
    with ht1:
        if not hallazgos:
            st.markdown('<div class="card" style="text-align:center;padding:36px;color:#6b7280">'
                        '<br>📋<br><br>Aún no hay hallazgos registrados.<br>'
                        'Usá la pestaña <b>➕ Nuevo Hallazgo</b> para registrar el primero.<br><br></div>',
                        unsafe_allow_html=True)
        else:
            # Métricas
            h1,h2,h3,h4 = st.columns(4)
            with h1: st.metric("NCM abiertas",  sum(1 for h in hallazgos if h["tipo"]=="No Conformidad Mayor" and h["estado"]!="verificado"))
            with h2: st.metric("NCn abiertas",  sum(1 for h in hallazgos if h["tipo"]=="No Conformidad Menor" and h["estado"]=="abierto"))
            with h3: st.metric("Obs / OPM",     sum(1 for h in hallazgos if h["tipo"] in ["Observación","Oportunidad de Mejora"]))
            with h4: st.metric("Verificados",   sum(1 for h in hallazgos if h["estado"]=="verificado"))

            # Filtros
            ff1, ff2, ff3 = st.columns([3,2,2])
            with ff1: f_tipo  = st.selectbox("Tipo", ["Todos"]+TIPOS_HALLAZGO, key="hf_tipo")
            with ff2: f_est   = st.selectbox("Estado", ["Todos"]+ESTADOS_HALLAZGO, key="hf_est")
            with ff3: f_clausula = st.text_input("Buscar cláusula", key="hf_claus", placeholder="ej: 7.5")

            items = hallazgos
            if f_tipo    != "Todos": items = [h for h in items if h["tipo"]==f_tipo]
            if f_est     != "Todos": items = [h for h in items if h["estado"]==f_est]
            if f_clausula.strip():   items = [h for h in items if f_clausula.strip() in h.get("clausula","")]

            if not items:
                st.info("No hay hallazgos con los filtros seleccionados.")
            else:
                for h in sorted(items, key=lambda x: x.get("fecha_auditoria",""), reverse=True):
                    card_cls = "hallazgo-card " + HCARD_CSS.get(h["tipo"],"hallazgo-obs")
                    st.markdown(
                        f'<div class="{card_cls}">'
                        f'<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px">'
                        f'{hallazgo_tipo_html(h["tipo"])} {hallazgo_estado_html(h["estado"])}'
                        f'<span style="font-size:.72rem;color:#9ca3af">📅 {h.get("fecha_auditoria","")} '
                        f'· 👤 {h.get("auditor","")}</span></div>'
                        f'<div style="font-size:.76rem;color:#6b7280;margin-bottom:6px">'
                        f'📌 Cláusula: <b>{h.get("clausula","—")}</b> &nbsp;·&nbsp; '
                        f'🏭 Proceso: <b>{h.get("proceso","—")}</b></div>'
                        f'<div style="font-weight:700;color:#1e293b;font-size:.9rem;margin-bottom:4px">'
                        f'{h.get("descripcion","")}</div>'
                        f'<div style="font-size:.81rem;color:#374151"><b>Evidencia:</b> '
                        f'{h.get("evidencia","—")}</div>'
                        + (f'<div style="font-size:.81rem;color:#374151;margin-top:4px">'
                           f'<b>Plan de acción:</b> {h.get("plan_accion","—")}'
                           f' &nbsp;·&nbsp; 👤 {h.get("responsable","—")}'
                           f' &nbsp;·&nbsp; 📆 {h.get("fecha_limite","—")}</div>'
                           if h.get("plan_accion") else "")
                        + f'</div>',
                        unsafe_allow_html=True,
                    )
                    hb1, hb2, hb3 = st.columns([2,2,2])
                    with hb1:
                        new_est_options = [e for e in ESTADOS_HALLAZGO if e != h["estado"]]
                        nxt = st.selectbox("Cambiar estado:", new_est_options,
                                           key=f"hest_{h['id']}", label_visibility="collapsed")
                        if st.button("Actualizar estado", key=f"hact_{h['id']}", use_container_width=True):
                            idx = next(i for i,x in enumerate(hallazgos) if x["id"]==h["id"])
                            hallazgos[idx]["estado"] = nxt
                            save_hallazgos(hallazgos); st.rerun()
                    with hb2:
                        if st.button("🤖 Asistencia IA", key=f"hai_{h['id']}", use_container_width=True):
                            st.session_state[f"hai_show_{h['id']}"] = True
                    with hb3:
                        if st.button("🗑️ Eliminar", key=f"hdel_{h['id']}", use_container_width=True):
                            hallazgos = [x for x in hallazgos if x["id"]!=h["id"]]
                            save_hallazgos(hallazgos); st.rerun()

                    if st.session_state.get(f"hai_show_{h['id']}"):
                        with st.spinner("Analizando hallazgo..."):
                            resp = analyze_hallazgo_redaccion(h["tipo"], h.get("clausula",""), h.get("descripcion",""))
                        st.session_state.pop(f"hai_show_{h['id']}", None)
                        with st.expander("🤖 Orientación del Asistente", expanded=True):
                            st.markdown(resp)

    # ── NUEVO HALLAZGO ────────────────────────────────────────────────────────
    with ht2:
        st.markdown('<div class="card-info" style="color:#1e293b!important;font-size:.84rem">'
                    'Complete los campos del hallazgo detectado. Podés pedir a la IA que sugiera '
                    'la <b>redacción formal</b> y el <b>plan de acción</b> correctivo.</div>',
                    unsafe_allow_html=True)

        with st.form("form_nuevo_hallazgo", clear_on_submit=True):
            nf1, nf2 = st.columns(2)
            with nf1:
                n_tipo     = st.selectbox("Tipo de hallazgo *", TIPOS_HALLAZGO)
                n_proceso  = st.text_input("Proceso / Área auditada *", placeholder="Ej: Gestión de Compras")
                n_clausula = st.text_input("Cláusula ISO relacionada *", placeholder="Ej: ISO 9001:2015 — 7.5.2")
                n_auditor  = st.text_input("Auditor", placeholder="Nombre del auditor")
            with nf2:
                n_fecha    = st.text_input("Fecha de auditoría", value=datetime.now().strftime("%Y-%m-%d"))
                n_resp     = st.text_input("Responsable de la acción", placeholder="Cargo o nombre")
                n_flimit   = st.text_input("Fecha límite de cierre", placeholder="YYYY-MM-DD")
                n_est_ini  = st.selectbox("Estado inicial", ESTADOS_HALLAZGO)

            n_descripcion = st.text_area("Descripción del hallazgo *",
                                         placeholder="Describa qué se observó, en qué no se cumple el requisito...",
                                         height=100)
            n_evidencia   = st.text_area("Evidencia objetiva observada",
                                         placeholder="Documentos revisados, registros inspeccionados, entrevistas...",
                                         height=80)
            n_plan        = st.text_area("Plan de acción correctiva/preventiva",
                                         placeholder="Dejá vacío y usá el botón 'Asistencia IA' para que Gemini lo genere...",
                                         height=80)
            n_obs         = st.text_area("Observaciones adicionales", height=68)

            submitted = st.form_submit_button("💾 Registrar Hallazgo", use_container_width=True,
                                              type="primary")
            if submitted:
                if not n_descripcion.strip() or not n_proceso.strip() or not n_clausula.strip():
                    st.error("Completá los campos obligatorios: tipo, proceso, cláusula y descripción.")
                else:
                    h_id = "H" + hashlib.md5(f"{n_fecha}{n_descripcion}".encode()).hexdigest()[:6].upper()
                    nuevo = {
                        "id": h_id, "tipo": n_tipo, "fecha_auditoria": n_fecha,
                        "auditor": n_auditor, "proceso": n_proceso, "clausula": n_clausula,
                        "descripcion": n_descripcion.strip(), "evidencia": n_evidencia.strip(),
                        "plan_accion": n_plan.strip(), "responsable": n_resp,
                        "fecha_limite": n_flimit, "estado": n_est_ini,
                        "observaciones": n_obs.strip(), "fecha_cierre": "", "verificacion": "",
                    }
                    hallazgos.append(nuevo); save_hallazgos(hallazgos)
                    st.success(f"✅ Hallazgo **{h_id}** registrado correctamente.")
                    st.rerun()

        # Asistencia IA para redacción
        st.markdown("---")
        st.markdown('<p class="sec-title">🤖 Asistencia IA para redacción de hallazgos</p>',
                    unsafe_allow_html=True)
        ai_tipo   = st.selectbox("Tipo de hallazgo", TIPOS_HALLAZGO, key="ai_h_tipo")
        ai_claus  = st.text_input("Cláusula", placeholder="Ej: 7.5.2", key="ai_h_claus")
        ai_desc   = st.text_area("Descripción libre (lo que observaste)", height=100, key="ai_h_desc",
                                  placeholder="Describí con tus palabras qué encontraste...")
        if st.button("🤖 Generar redacción formal + plan de acción", use_container_width=True,
                     key="ai_h_gen", type="primary"):
            if not ai_desc.strip():
                st.warning("Ingresá una descripción del hallazgo.")
            else:
                with st.spinner("Generando redacción formal..."):
                    resp = analyze_hallazgo_redaccion(ai_tipo, ai_claus, ai_desc)
                st.markdown(
                    f'<div class="card" style="margin-top:12px">{resp}</div>',
                    unsafe_allow_html=True,
                )

    # ── IMPORTAR DESDE INFORME ────────────────────────────────────────────────
    with ht3:
        st.markdown(
            '<div class="card-info" style="color:#1e293b!important;font-size:.84rem">'
            '📤 Subí un <b>informe o plan de auditoría</b> (PDF, Word, TXT) y la IA extraerá '
            'automáticamente todos los hallazgos detectados para que puedas revisarlos e '
            'importarlos directamente al registro.</div>',
            unsafe_allow_html=True,
        )

        imp_file = st.file_uploader(
            "Seleccioná el informe de auditoría",
            type=["pdf","docx","txt"],
            key="imp_audit_file",
            help="Informes de auditoría, planes, actas — PDF, Word o TXT",
        )

        if imp_file:
            ic1, ic2 = st.columns([2, 3])
            with ic1:
                extract_btn = st.button(
                    "🤖 Extraer Hallazgos con IA",
                    type="primary",
                    use_container_width=True,
                    key="btn_extract_hllzg",
                )
            with ic2:
                st.markdown(
                    '<div class="card" style="font-size:.82rem;color:#374151;padding:10px 14px">'
                    '💡 Funciona mejor con informes que listen hallazgos, NC y observaciones '
                    'de forma explícita. También analiza planes y actas de reunión.</div>',
                    unsafe_allow_html=True,
                )

            if extract_btn:
                with st.spinner(f"Analizando «{imp_file.name}» con Gemini..."):
                    imp_text = extract_text(imp_file.getvalue(), imp_file.name)
                if not imp_text.strip():
                    st.error("No se pudo extraer texto del archivo. "
                             "Si es PDF escaneado, aplicá OCR primero.")
                else:
                    with st.spinner("Identificando hallazgos..."):
                        extracted = extract_hallazgos_from_report(imp_text, imp_file.name)
                    st.session_state["extracted_hallazgos"] = extracted
                    st.session_state["import_report_name"]  = imp_file.name
                    if not extracted:
                        st.warning(
                            "⚠️ No se encontraron hallazgos en el documento. "
                            "Verificá que sea un informe de auditoría con hallazgos explícitos."
                        )
                    else:
                        st.success(f"✅ Se encontraron **{len(extracted)}** hallazgo(s). "
                                   f"Revisalos abajo y seleccioná los que querés importar.")

        # ── Listado de hallazgos extraídos ────────────────────────────────────
        extracted_list = st.session_state.get("extracted_hallazgos", [])
        report_src     = st.session_state.get("import_report_name", "")

        if extracted_list:
            st.markdown("---")
            st.markdown(
                f'<p class="sec-title">📋 {len(extracted_list)} hallazgo(s) extraído(s) '
                f'de «{report_src}»</p>',
                unsafe_allow_html=True,
            )
            st.markdown(
                '<div class="card-warn" style="color:#1e293b!important;font-size:.83rem">'
                '⚠️ Revisá cada hallazgo antes de importar. Podés desmarcar los que no '
                'correspondan o ya existan en el registro.</div>',
                unsafe_allow_html=True,
            )

            # Botones de selección masiva
            ms1, ms2, _ = st.columns([1.2, 1.5, 4])
            with ms1:
                if st.button("☑️ Todos", key="imp_sel_all", use_container_width=True):
                    for idx in range(len(extracted_list)):
                        st.session_state[f"imp_chk_{idx}"] = True
            with ms2:
                if st.button("☐ Ninguno", key="imp_desel_all", use_container_width=True):
                    for idx in range(len(extracted_list)):
                        st.session_state[f"imp_chk_{idx}"] = False

            # Tarjetas con checkbox
            selected_to_import = []
            for i, h in enumerate(extracted_list):
                tipo_h = h.get("tipo","Observación")
                card_cls = "hallazgo-card " + HCARD_CSS.get(tipo_h, "hallazgo-obs")
                chk_col, card_col = st.columns([0.4, 11])
                with chk_col:
                    checked = st.checkbox(
                        "", value=st.session_state.get(f"imp_chk_{i}", True),
                        key=f"imp_chk_{i}",
                        label_visibility="collapsed",
                    )
                with card_col:
                    st.markdown(
                        f'<div class="{card_cls}" style="margin-bottom:4px">'
                        f'<div style="display:flex;justify-content:space-between;'
                        f'align-items:center;margin-bottom:5px">'
                        f'{hallazgo_tipo_html(tipo_h)}'
                        f'<span style="font-size:.72rem;color:#9ca3af">'
                        f'📌 {h.get("clausula","—")} &nbsp;·&nbsp; '
                        f'🏭 {h.get("proceso","—")}'
                        + (f' &nbsp;·&nbsp; 👤 {h["auditor"]}' if h.get("auditor") else "")
                        + f'</span></div>'
                        f'<div style="font-size:.88rem;font-weight:600;color:#1e293b;'
                        f'margin-bottom:4px">{h.get("descripcion","")}</div>'
                        + (f'<div style="font-size:.79rem;color:#374151">'
                           f'<b>Evidencia:</b> {h["evidencia"]}</div>'
                           if h.get("evidencia") else "")
                        + (f'<div style="font-size:.79rem;color:#374151;margin-top:3px">'
                           f'<b>Plan de acción:</b> {h["plan_accion"]}'
                           + (f' &nbsp;·&nbsp; 👤 {h["responsable"]}' if h.get("responsable") else "")
                           + (f' &nbsp;·&nbsp; 📆 {h["fecha_limite"]}' if h.get("fecha_limite") else "")
                           + '</div>'
                           if h.get("plan_accion") else "")
                        + '</div>',
                        unsafe_allow_html=True,
                    )
                if checked:
                    selected_to_import.append(i)

            # Botón importar
            st.markdown("---")
            n_sel = len(selected_to_import)
            btn_imp_col, info_imp_col = st.columns([2, 3])
            with btn_imp_col:
                import_disabled = (n_sel == 0)
                if st.button(
                    f"⬆️ Importar {n_sel} hallazgo(s) al registro",
                    type="primary",
                    use_container_width=True,
                    key="btn_import_hllzg",
                    disabled=import_disabled,
                ):
                    existing_hallazgos = load_hallazgos()
                    fecha_hoy = datetime.now().strftime("%Y-%m-%d")
                    n_importados = 0
                    for idx in selected_to_import:
                        h = extracted_list[idx]
                        h_id = ("H" + hashlib.md5(
                            f"{fecha_hoy}{h.get('descripcion','')}{idx}{report_src}".encode()
                        ).hexdigest()[:6].upper())
                        # Evitar duplicados por id
                        if any(x["id"] == h_id for x in existing_hallazgos):
                            continue
                        nuevo = {
                            "id":            h_id,
                            "tipo":          h.get("tipo","Observación"),
                            "fecha_auditoria": h.get("fecha_auditoria", fecha_hoy),
                            "auditor":       h.get("auditor",""),
                            "proceso":       h.get("proceso",""),
                            "clausula":      h.get("clausula",""),
                            "descripcion":   h.get("descripcion","").strip(),
                            "evidencia":     h.get("evidencia","").strip(),
                            "plan_accion":   h.get("plan_accion","").strip(),
                            "responsable":   h.get("responsable",""),
                            "fecha_limite":  h.get("fecha_limite",""),
                            "estado":        "abierto",
                            "observaciones": f"Importado desde: {report_src}",
                            "fecha_cierre":  "",
                            "verificacion":  "",
                        }
                        existing_hallazgos.append(nuevo)
                        n_importados += 1
                    save_hallazgos(existing_hallazgos)
                    # Limpiar estado
                    st.session_state["extracted_hallazgos"] = []
                    for idx in range(len(extracted_list)):
                        st.session_state.pop(f"imp_chk_{idx}", None)
                    st.success(
                        f"✅ **{n_importados}** hallazgo(s) importado(s) correctamente. "
                        f"Vérlos en la pestaña 📋 Registro de Hallazgos."
                    )
                    st.rerun()

            with info_imp_col:
                if n_sel > 0:
                    st.markdown(
                        f'<div class="card-info" style="color:#1e293b!important;font-size:.83rem">'
                        f'Se importarán <b>{n_sel}</b> hallazgo(s) con estado inicial '
                        f'<b>abierto</b>. Desde el Registro podrás editar estado, '
                        f'plan de acción y responsable de cada uno.</div>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        '<div class="card-warn" style="color:#1e293b!important;font-size:.83rem">'
                        'Seleccioná al menos un hallazgo para importar.</div>',
                        unsafe_allow_html=True,
                    )

            # Botón limpiar
            if st.button("🗑️ Limpiar resultados", key="btn_clear_extracted"):
                st.session_state["extracted_hallazgos"] = []
                for idx in range(len(extracted_list)):
                    st.session_state.pop(f"imp_chk_{idx}", None)
                st.rerun()


# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
def render_sidebar(lista_maestra: list, incongruencias: dict):
    with st.sidebar:
        st.markdown("### 🏛️ SGI Asistente v3")
        st.markdown("---")
        has_gemini = bool(os.environ.get("GEMINI_API_KEY"))
        has_groq   = bool(os.environ.get("GROQ_API_KEY"))
        st.markdown(f"{'🟢' if has_gemini else '🔴'} &nbsp;**Gemini** · "
                    f"{'Activa' if has_gemini else 'Sin clave'}",
                    unsafe_allow_html=True)
        st.markdown(f"{'🟢' if has_groq else '⚪'} &nbsp;**Groq (Llama 3.3)** · "
                    f"{'Activa' if has_groq else 'Sin configurar — fallback gratuito'}",
                    unsafe_allow_html=True)
        if not has_groq:
            st.markdown('<div style="font-size:.71rem;color:rgba(255,255,255,.55);'
                        'margin-top:2px">👉 groq.com → clave gratis → agregar<br>'
                        'GROQ_API_KEY=xxx al archivo .env</div>',
                        unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("### 📊 Estado del Sistema")
        activos  = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
        obsoletos= [d for d in lista_maestra if d.get("estado","activo")=="obsoleto"]
        abiertas = sum(1 for v in incongruencias.values() if v.get("estado")=="abierta")
        try:    n_chunks = get_collection().count()
        except: n_chunks = 0
        for val, lbl in [(len(activos),"Docs Activos"),(len(obsoletos),"Obsoletos"),
                         (abiertas,"Inc. Abiertas"),(n_chunks,"Chunks RAG")]:
            st.markdown(f'<div class="stat-card"><div class="stat-val">{val}</div>'
                        f'<div class="stat-lbl">{lbl}</div></div>', unsafe_allow_html=True)
        if activos:
            st.markdown("---")
            st.markdown("### 🏷️ Tipos")
            tipos = {}
            for d in activos: tipos[d.get("tipo","Otro")] = tipos.get(d.get("tipo","Otro"),0)+1
            for t,cnt in sorted(tipos.items(),key=lambda x:-x[1]):
                st.markdown(f"&nbsp;&nbsp;{tipo_badge(t)}&nbsp; ×{cnt}", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown('<div class="foot-info">🤖 Gemini 2.5 Flash<br>🔍 ChromaDB local<br>'
                    '📐 all-MiniLM-L6-v2<br>📄 python-docx<br><br>'
                    '📁 lista_maestra.json<br>📁 incongruencias.json<br>'
                    '📁 revisiones.json<br>📁 ./chroma_db/</div>',
                    unsafe_allow_html=True)


# ─── TAB: DOCUMENTOS ──────────────────────────────────────────────────────────
def tab_documentos(lista_maestra: list, analisis_cache: dict, incongruencias: dict):
    st.markdown('<p class="sec-title">📂 Incorporar Documento al SGI</p>', unsafe_allow_html=True)
    st.markdown('<div class="card-info" style="color:#1e293b!important">Suba un documento para indexarlo, '
                'clasificarlo según ISO 9001/39001 e incorporarlo a la Lista Maestra. '
                'El análisis se calcula <b>una sola vez</b> y queda guardado localmente.</div>',
                unsafe_allow_html=True)
    uploaded = st.file_uploader("Seleccione o arrastre un documento",
                                type=["pdf","docx","xlsx","html","htm","txt"],
                                key="main_uploader",
                                help="PDF, Word, Excel, HTML (indicadores), TXT")
    if uploaded is None: return lista_maestra, analisis_cache, incongruencias

    file_bytes = uploaded.getvalue()
    file_hash  = hashlib.md5(file_bytes).hexdigest()[:8]
    already    = any(d.get("hash")==file_hash for d in lista_maestra)
    if already:
        doc = next((d for d in lista_maestra if d.get("hash")==file_hash), {})
        st.markdown(f'<div class="card-ok">✅ <b>{uploaded.name}</b> ya está en la Lista Maestra como '
                    f'{tipo_badge(doc.get("tipo",""))}. Para analizar cambios use <b>🔄 Revisiones</b>.</div>',
                    unsafe_allow_html=True)
        return lista_maestra, analisis_cache, incongruencias

    prog = st.progress(0, text="Iniciando...")
    with st.status(f"⚙️ Procesando **{uploaded.name}**...", expanded=True) as status:
        st.write("📄 Extrayendo texto...")
        prog.progress(15); text = extract_text(file_bytes, uploaded.name)
        if not text.strip():
            prog.empty(); status.update(label="❌ Sin texto extraíble.", state="error")
            st.error("No se pudo extraer texto. Si es PDF escaneado aplique OCR primero.")
            return lista_maestra, analisis_cache, incongruencias
        st.write("🔵 Clasificando según ISO 9001 e ISO 39001...")
        prog.progress(35); classification = classify_document(text, uploaded.name)
        st.write("🔍 Analizando incongruencias y sugerencias (una sola vez)...")
        prog.progress(58); deep = analyze_document_deep(text, uploaded.name)
        st.write("🧠 Vectorizando e indexando en ChromaDB...")
        prog.progress(78); index_document(file_bytes, uploaded.name, text)
        st.write("💾 Guardando en Lista Maestra y tracker de incongruencias...")
        prog.progress(92)
        entry = {"hash":file_hash,"nombre":uploaded.name,"tipo":classification.get("tipo","Otro"),
                 "iso9001":classification.get("iso9001","No aplica"),
                 "iso39001":classification.get("iso39001","No aplica"),
                 "estado":"activo","fecha_ingreso":datetime.now().strftime("%Y-%m-%d")}
        lista_maestra.append(entry); save_json(LISTA_MAESTRA_PATH, lista_maestra)
        analisis_cache[uploaded.name] = {**deep, "_texto": text[:TEXT_PREVIEW_LEN]}
        save_json(ANALISIS_PATH, analisis_cache)
        # Agregar incongruencias al tracker (sin duplicar las que ya existen en otros docs)
        inc_nuevas = 0; inc_descartadas = 0
        for inc_txt in deep.get("incongruencias",[]):
            if inc_txt and not any(p in inc_txt for p in ["⚠️","📌","Error","Configure"]):
                if _incongruencia_ya_existe(inc_txt, incongruencias, uploaded.name):
                    inc_descartadas += 1
                else:
                    incongruencias = add_incongruencia_from_analysis(uploaded.name, inc_txt, incongruencias)
                    inc_nuevas += 1
        if inc_descartadas:
            st.info(f"ℹ️ {inc_descartadas} incongruencia(s) omitida(s) por ya existir en otros documentos.")
        save_incongruencias(incongruencias)
        prog.progress(100); status.update(label=f"✅ **{uploaded.name}** incorporado.", state="complete")
    prog.empty(); st.rerun()
    return lista_maestra, analisis_cache, incongruencias


# ─── TAB: REVISIONES ──────────────────────────────────────────────────────────
def tab_revisiones(lista_maestra: list, analisis_cache: dict):
    revisiones = load_json(REVISIONES_PATH, {})
    sugerencias_nue = load_json(SUGERENCIAS_NUE_PATH, {})
    activos = [d for d in lista_maestra if d.get("estado","activo")=="activo"]

    st.markdown('<p class="sec-title">🔄 Análisis de Revisiones y Nuevos Documentos</p>', unsafe_allow_html=True)
    modo = st.radio("Modo:", ["🔄  Nueva revisión de documento existente",
                               "➕  Documento nuevo a incorporar al SGI"],
                    horizontal=True, label_visibility="collapsed")
    st.markdown("<br>", unsafe_allow_html=True)

    # ── MODO A: REVISIÓN ──────────────────────────────────────────────────────
    if "revisión" in modo or "revision" in modo.lower():
        st.markdown('<div class="card-info" style="color:#1e293b!important">'
                    'Suba la <b>nueva versión</b> del documento. El sistema recuperará el texto original '
                    'desde ChromaDB y realizará una comparación <b>cláusula por cláusula</b> contra ISO 9001:2015 '
                    'e ISO 39001:2015 para emitir un veredicto normativo.</div>', unsafe_allow_html=True)
        if not activos:
            st.warning("No hay documentos activos. Suba documentos en la pestaña 📂 Documentos.")
            return
        ca, cb = st.columns(2, gap="medium")
        with ca:
            doc_original = st.selectbox("¿Cuál documento actualiza?", [d["nombre"] for d in activos])
        with cb:
            rev_file = st.file_uploader("Suba la nueva versión", type=["pdf","docx","xlsx"], key="rev_uploader")

        if rev_file and doc_original:
            if st.button("🔍 Analizar Cambios", type="primary", use_container_width=True):
                with st.spinner("Recuperando texto original y comparando cláusula por cláusula..."):
                    old_text = (analisis_cache.get(doc_original,{}).get("_texto","")
                                or get_document_text_from_chroma(doc_original))
                    if not old_text:
                        st.error(f"No se encontró el texto de **{doc_original}**. Vuelva a subirlo desde 📂 Documentos.")
                        return
                    new_text = extract_text(rev_file.getvalue(), rev_file.name)
                    if not new_text.strip():
                        st.error("No se pudo extraer texto de la nueva versión."); return
                    result = analyze_revision(old_text, new_text, doc_original, rev_file.name)
                    record = {"fecha":datetime.now().strftime("%Y-%m-%d %H:%M"),
                              "nueva_version":rev_file.name, **result}
                    revisiones.setdefault(doc_original,[]).insert(0,record)
                    save_json(REVISIONES_PATH, revisiones)
                    st.session_state["last_revision"] = record
                    st.session_state["last_rev_original"] = doc_original

        result = st.session_state.get("last_revision")
        if result:
            st.markdown("---")
            st.markdown(f'<div style="text-align:center;margin:12px 0">{verdict_html(result.get("veredicto",""))}</div>',
                        unsafe_allow_html=True)
            st.markdown(f'<div class="card-info" style="color:#1e293b!important"><b>Resumen:</b> '
                        f'{result.get("resumen_ejecutivo","")}</div>', unsafe_allow_html=True)

            # Tabla de cláusulas verificadas
            clausulas = result.get("clausulas_verificadas", [])
            if clausulas:
                st.markdown("#### 📋 Verificación por Cláusula ISO")
                icons = {"OK":"✅","OBSERVACION":"⚠️","NO_APLICA":"➖","ERROR":"❌"}
                tbl_data = [{"Cláusula":c.get("clausula",""),
                             "Estado": icons.get(c.get("resultado",""),"?") + " " + c.get("resultado",""),
                             "Detalle": c.get("detalle","")} for c in clausulas]
                st.dataframe(tbl_data, use_container_width=True, hide_index=True,
                             column_config={"Cláusula":st.column_config.TextColumn(width="medium"),
                                            "Estado":st.column_config.TextColumn(width="small"),
                                            "Detalle":st.column_config.TextColumn(width="large")})

            c1, c2 = st.columns(2)
            with c1:
                with st.expander("📝 Cambios Detectados", expanded=True):
                    for i,c in enumerate(result.get("cambios_detectados",[]),1):
                        st.markdown(step_html(i,c), unsafe_allow_html=True)
                with st.expander("⚠️ Riesgos", expanded=True):
                    for i,r in enumerate(result.get("riesgos",[]),1):
                        st.markdown(f'<div class="card-warn" style="color:#1e293b!important"><b>{i}.</b> {r}</div>',
                                    unsafe_allow_html=True)
            with c2:
                with st.expander("✅ Recomendaciones", expanded=True):
                    for i,r in enumerate(result.get("recomendaciones",[]),1):
                        st.markdown(f'<div class="card-ok" style="color:#1e293b!important"><b>{i}.</b> {r}</div>',
                                    unsafe_allow_html=True)

        hist = revisiones.get(doc_original if activos else "", [])
        if hist:
            st.markdown("---")
            st.markdown('<p class="rpanel-lbl">📅 Historial de revisiones</p>', unsafe_allow_html=True)
            for rec in hist[:5]:
                v = rec.get("veredicto",""); ico = "✅" if "APROBADO" in v else ("⚠️" if "OBSERV" in v else "❌")
                with st.expander(f"{ico} {rec.get('fecha','')} → {rec.get('nueva_version','')} · {v}"):
                    st.markdown(rec.get("resumen_ejecutivo",""))
                    for c in rec.get("cambios_detectados",[]): st.markdown(f"- {c}")

    # ── MODO B: DOCUMENTO NUEVO ───────────────────────────────────────────────
    else:
        st.markdown('<div class="card-info" style="color:#1e293b!important">Suba un documento que aún '
                    '<b>no está en el SGI</b>. Gemini analizará dónde encaja, qué documentos se relacionan '
                    'y qué vacíos viene a cubrir.</div>', unsafe_allow_html=True)
        sug_file = st.file_uploader("Suba el nuevo documento", type=["pdf","docx","xlsx"], key="sug_uploader")
        if sug_file:
            sug = sugerencias_nue.get(sug_file.name)
            if sug: st.info(f"ℹ️ Se encontró análisis previo para **{sug_file.name}** (caché).")
            if not sug or st.button("🔄 Regenerar análisis"):
                if st.button("🧭 Sugerir Ubicación en el SGI", type="primary", use_container_width=True):
                    with st.spinner("Analizando y comparando con el SGI existente..."):
                        new_text = extract_text(sug_file.getvalue(), sug_file.name)
                        if not new_text.strip(): st.error("No se pudo extraer texto."); return
                        sug = suggest_placement(new_text, sug_file.name, activos)
                        sugerencias_nue[sug_file.name] = {"fecha":datetime.now().strftime("%Y-%m-%d %H:%M"),**sug}
                        save_json(SUGERENCIAS_NUE_PATH, sugerencias_nue)
                        st.session_state["last_sugerencia"] = sug
            sug = sug or st.session_state.get("last_sugerencia")
            if sug:
                st.markdown("---")
                cx, cy = st.columns([2,1])
                with cx:
                    st.markdown(f'<div class="card"><b>Tipo:</b> {tipo_badge(sug.get("tipo_sugerido",""))}<br><br>'
                                f'<b>📍 Ubicación en el SGI:</b><br>{sug.get("ubicacion_sgi","")}</div>',
                                unsafe_allow_html=True)
                with cy:
                    st.markdown(f'<div class="card-info" style="color:#1e293b!important"><b>🔵 ISO 9001:2015</b><br>'
                                f'{sug.get("iso9001_sugerido","")}<br><br><b>🟢 ISO 39001:2015</b><br>'
                                f'{sug.get("iso39001_sugerido","")}</div>', unsafe_allow_html=True)
                c1,c2 = st.columns(2)
                with c1:
                    with st.expander("🔗 Documentos Relacionados", expanded=True):
                        for d in sug.get("documentos_relacionados",[]):
                            st.markdown(step_html("📄",d), unsafe_allow_html=True)
                    with st.expander("🟡 Vacíos que Cubre", expanded=True):
                        for i,v in enumerate(sug.get("vacios_que_cubre",[]),1):
                            st.markdown(f'<div class="card-warn" style="color:#1e293b!important"><b>{i}.</b> {v}</div>',
                                        unsafe_allow_html=True)
                with c2:
                    with st.expander("💡 Justificación", expanded=True):
                        st.markdown(f'<div class="card" style="color:#1e293b!important">{sug.get("razon","")}</div>',
                                    unsafe_allow_html=True)
                    st.markdown('<div class="card-ok" style="color:#1e293b!important">Confirmada la ubicación, '
                                'suba este documento desde <b>📂 Documentos</b> para incorporarlo oficialmente.</div>',
                                unsafe_allow_html=True)


# ─── TAB: INCONGRUENCIAS ──────────────────────────────────────────────────────
def tab_incongruencias(lista_maestra: list, analisis_cache: dict, incongruencias: dict):
    st.markdown('<p class="sec-title">⚠️ Gestión de Incongruencias</p>', unsafe_allow_html=True)

    activos = [d["nombre"] for d in lista_maestra if d.get("estado","activo")=="activo"]
    all_items = [v for v in incongruencias.values() if v.get("estado") != "eliminada"]

    # ── Filtros ────────────────────────────────────────────────────────────────
    f1, f2, f3 = st.columns([3,2,2])
    with f1:
        doc_filter = st.selectbox("📂 Filtrar por documento",
                                  ["Todos los documentos"] + sorted({v["documento"] for v in all_items}),
                                  key="inc_doc_filter")
    with f2:
        estado_filter = st.selectbox("Estado", ["Todas","abierta","resuelta"], key="inc_estado_filter")
    with f3:
        st.markdown("<br>", unsafe_allow_html=True)
        n_ab = sum(1 for v in all_items if v.get("estado")=="abierta")
        n_rs = sum(1 for v in all_items if v.get("estado")=="resuelta")
        st.markdown(f'<span class="status-pill st-abierta">🔴 {n_ab} abiertas</span>&nbsp;&nbsp;'
                    f'<span class="status-pill st-resuelta">✅ {n_rs} resueltas</span>',
                    unsafe_allow_html=True)

    # Aplicar filtros
    filtered = all_items
    if doc_filter != "Todos los documentos":
        filtered = [v for v in filtered if v["documento"] == doc_filter]
    if estado_filter != "Todas":
        filtered = [v for v in filtered if v.get("estado") == estado_filter]

    if not filtered:
        st.info("No hay incongruencias que coincidan con los filtros seleccionados.")
        return

    st.markdown("---")

    # ── Layout: lista | chat ────────────────────────────────────────────────
    col_list, col_chat = st.columns([4, 6], gap="large")

    selected_id = st.session_state.get("inc_selected")

    with col_list:
        st.markdown(f'<p class="rpanel-lbl">📋 {len(filtered)} incongruencia(s)</p>', unsafe_allow_html=True)
        for item in sorted(filtered, key=lambda x: x.get("fecha",""), reverse=True):
            is_active = (item["id"] == selected_id)
            card_class = "inc-card inc-card-active" if is_active else "inc-card"
            st.markdown(
                f'<div class="{card_class}">'
                f'<div style="font-size:.72rem;color:#6b7280;margin-bottom:4px">'
                f'📄 {item["documento"]}&nbsp;&nbsp;{status_pill(item.get("estado","abierta"))}</div>'
                f'<div style="font-size:.87rem;color:#1e293b;font-weight:500;line-height:1.4">'
                f'{item["texto"][:140]}{"..." if len(item["texto"])>140 else ""}</div>'
                f'<div style="font-size:.71rem;color:#9ca3af;margin-top:4px">🗓 {item.get("fecha","")}'
                f' · 💬 {len(item.get("chat",[]))} mensajes</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            bc1, bc2, bc3 = st.columns([2,2,1])
            with bc1:
                if st.button("💬 Chatear", key=f"sel_{item['id']}", use_container_width=True):
                    st.session_state["inc_selected"] = item["id"]
                    st.rerun()
            with bc2:
                new_estado = "resuelta" if item.get("estado")=="abierta" else "abierta"
                btn_lbl = "✅ Resolver" if item.get("estado")=="abierta" else "🔄 Reabrir"
                if st.button(btn_lbl, key=f"res_{item['id']}", use_container_width=True):
                    incongruencias[item["id"]]["estado"] = new_estado
                    save_incongruencias(incongruencias)
                    st.rerun()
            with bc3:
                if st.button("🗑️", key=f"del_{item['id']}", use_container_width=True,
                             help="Eliminar incongruencia"):
                    incongruencias[item["id"]]["estado"] = "eliminada"
                    save_incongruencias(incongruencias)
                    if st.session_state.get("inc_selected") == item["id"]:
                        st.session_state["inc_selected"] = None
                    st.rerun()

    # ── Panel de chat ──────────────────────────────────────────────────────────
    with col_chat:
        if not selected_id or selected_id not in incongruencias:
            st.markdown('<div class="card" style="text-align:center;padding:40px;color:#6b7280;">'
                        '<br>💬<br><br>Seleccioná una incongruencia<br>de la lista para chatear con la IA<br><br></div>',
                        unsafe_allow_html=True)
            return

        inc = incongruencias[selected_id]
        st.markdown(
            f'<div class="card-warn" style="color:#1e293b!important"><b>Incongruencia seleccionada:</b><br>'
            f'{inc["texto"]}<br><br>'
            f'<small>📄 {inc["documento"]} &nbsp;·&nbsp; {status_pill(inc.get("estado","abierta"))}</small></div>',
            unsafe_allow_html=True,
        )

        # Botones de acción rápida
        qa1, qa2, qa3 = st.columns(3)
        with qa1:
            if st.button("📋 ¿Cómo corregirla?", use_container_width=True, key="qa_corr"):
                st.session_state["inc_quick_q"] = "¿Cómo puedo corregir esta incongruencia? Dame el texto exacto a insertar en el documento."
        with qa2:
            if st.button("📌 ¿Qué cláusula aplica?", use_container_width=True, key="qa_claus"):
                st.session_state["inc_quick_q"] = "¿Cuál es la cláusula exacta de la norma ISO que se incumple y qué dice literalmente?"
        with qa3:
            if st.button("📈 ¿Qué indicador proponer?", use_container_width=True, key="qa_kpi"):
                st.session_state["inc_quick_q"] = "¿Qué indicador de desempeño (KPI) podría implementarse para controlar esta incongruencia?"

        # Historial de chat
        chat_box = st.container(height=480)
        with chat_box:
            if not inc.get("chat"):
                st.markdown(
                    f'<div class="card-info" style="color:#1e293b!important">'
                    f'🤖 <b>Asistente listo.</b> Estoy aquí para ayudarte a resolver esta incongruencia. '
                    f'Podés preguntarme cómo corregirla, qué cláusula aplica, qué texto insertar en el documento, '
                    f'o qué indicadores implementar.</div>',
                    unsafe_allow_html=True,
                )
            for msg in inc.get("chat", []):
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

        # Input de pregunta (o pregunta rápida)
        quick_q = st.session_state.pop("inc_quick_q", None)
        user_input = st.chat_input("Preguntá sobre esta incongruencia...", key=f"chat_inc_{selected_id}")
        question = quick_q or user_input

        if question:
            incongruencias[selected_id]["chat"].append({"role":"user","content":question})
            with chat_box:
                with st.chat_message("user"): st.markdown(question)
                with st.chat_message("assistant"):
                    with st.spinner("Analizando..."):
                        answer = chat_incongruencia(inc["texto"], inc["documento"],
                                                    inc.get("chat",[]), question, analisis_cache)
                    st.markdown(answer)
            incongruencias[selected_id]["chat"].append({"role":"assistant","content":answer})
            save_incongruencias(incongruencias)
            st.rerun()


# ─── TAB: REPOSITORIO ────────────────────────────────────────────────────────
def tab_repositorio(lista_maestra: list, analisis_cache: dict):
    st.markdown('<p class="sec-title">🗄️ Repositorio de Documentos</p>', unsafe_allow_html=True)

    activos   = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
    obsoletos = [d for d in lista_maestra if d.get("estado","activo")=="obsoleto"]

    sub1, sub2 = st.tabs(["✅  Documentos Activos","🗂️  Documentos Obsoletos"])

    with sub1:
        if not activos:
            st.info("No hay documentos activos. Suba documentos desde la pestaña 📂 Documentos.")
        else:
            st.markdown(f'<p class="rpanel-lbl">{len(activos)} documento(s) activo(s) en el SGI</p>',
                        unsafe_allow_html=True)
            for doc in activos:
                analysis = analisis_cache.get(doc["nombre"],{})
                with st.container():
                    st.markdown(
                        f'<div class="repo-card"><div style="display:flex;justify-content:space-between;'
                        f'align-items:flex-start;flex-wrap:wrap;gap:8px">'
                        f'<div><span style="font-weight:700;font-size:.95rem;color:#1a237e">{doc["nombre"]}</span>'
                        f'&nbsp;&nbsp;{tipo_badge(doc.get("tipo",""))}</div>'
                        f'<span style="font-size:.72rem;color:#6b7280">📅 {doc.get("fecha_ingreso","")}</span>'
                        f'</div>'
                        f'<div style="margin-top:8px;font-size:.82rem;color:#374151">'
                        f'<b>ISO 9001:</b> {doc.get("iso9001","")[:80]}{"..." if len(doc.get("iso9001",""))>80 else ""}<br>'
                        f'<b>ISO 39001:</b> {doc.get("iso39001","")[:80]}{"..." if len(doc.get("iso39001",""))>80 else ""}'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )
                    rc1, rc2, rc3 = st.columns([3,2,2])
                    with rc1:
                        n_inc = sum(1 for v in load_json(INCONGRUENCIAS_PATH,{}).values()
                                    if v.get("documento")==doc["nombre"] and v.get("estado")=="abierta")
                        if n_inc: st.markdown(f'<span class="status-pill st-abierta">⚠️ {n_inc} inc. abierta(s)</span>',
                                              unsafe_allow_html=True)
                    with rc2:
                        if st.button("👁️ Ver contenido", key=f"view_{doc['hash']}", use_container_width=True):
                            current = st.session_state.get("repo_view_doc")
                            st.session_state["repo_view_doc"] = None if current == doc["nombre"] else doc["nombre"]
                            st.rerun()
                    with rc3:
                        if st.button("📦 Archivar como Obsoleto", key=f"arch_{doc['hash']}", use_container_width=True):
                            for i,d in enumerate(lista_maestra):
                                if d.get("hash")==doc["hash"]:
                                    lista_maestra[i]["estado"] = "obsoleto"
                                    lista_maestra[i]["fecha_obsolescencia"] = datetime.now().strftime("%Y-%m-%d")
                                    break
                            save_json(LISTA_MAESTRA_PATH, lista_maestra)
                            st.success(f"**{doc['nombre']}** archivado como obsoleto.")
                            st.rerun()
                    # Visor de texto del documento
                    if st.session_state.get("repo_view_doc") == doc["nombre"]:
                        doc_texto = (analisis_cache.get(doc["nombre"], {}).get("_texto", "")
                                     or get_document_text_from_chroma(doc["nombre"]))
                        if doc_texto:
                            st.markdown(
                                f'<div class="card" style="margin-top:6px">'
                                f'<div style="font-size:.75rem;font-weight:700;color:#1565c0;'
                                f'margin-bottom:8px;text-transform:uppercase;letter-spacing:.5px">'
                                f'👁️ Contenido extraído — {doc["nombre"]}</div>'
                                f'<pre style="white-space:pre-wrap;font-size:.8rem;color:#1e293b;'
                                f'max-height:360px;overflow-y:auto;background:#f8fafc;'
                                f'padding:12px;border-radius:8px;border:1px solid #e2e8f0">'
                                f'{doc_texto[:8000]}{"\\n\\n[... texto truncado — primeros 8000 caracteres ...]" if len(doc_texto)>8000 else ""}'
                                f'</pre></div>',
                                unsafe_allow_html=True,
                            )
                        else:
                            st.info("No se encontró texto guardado para este documento. "
                                    "Volvé a subirlo desde 📂 Documentos para regenerar el análisis.")

                    # Análisis colapsable
                    with st.expander(f"🔍 Ver análisis de {doc['nombre'][:40]}...", expanded=False):
                        incongruencias_doc = [t for t in analysis.get("incongruencias",[])
                                              if not any(p in t for p in ["⚠️","Configure","Error"])]
                        sugerencias_doc    = [t for t in analysis.get("sugerencias",[])
                                              if not any(p in t for p in ["📌","Configure","Error"])]
                        if incongruencias_doc:
                            st.markdown("**⚠️ Incongruencias:**")
                            for t in incongruencias_doc:
                                st.markdown(f'<div class="card-warn" style="color:#1e293b!important;font-size:.84rem">{t}</div>',
                                            unsafe_allow_html=True)
                        if sugerencias_doc:
                            st.markdown("**📈 Sugerencias:**")
                            for t in sugerencias_doc:
                                st.markdown(f'<div class="card-ok" style="color:#1e293b!important;font-size:.84rem">{t}</div>',
                                            unsafe_allow_html=True)
                    st.markdown("<hr style='margin:8px 0;border-color:#e5e7eb!important'>", unsafe_allow_html=True)

    with sub2:
        if not obsoletos:
            st.info("No hay documentos archivados como obsoletos.")
        else:
            st.markdown(
                '<div style="background:#7f1d1d;color:white;padding:12px 18px;border-radius:10px;margin-bottom:16px;">'
                '⛔ <b>ÁREA DE DOCUMENTOS OBSOLETOS</b> — Estos documentos han sido reemplazados por versiones más recientes. '
                'No deben utilizarse como referencia vigente. Se conservan solo como registro histórico.</div>',
                unsafe_allow_html=True,
            )
            for doc in obsoletos:
                st.markdown(
                    f'<div class="repo-card" style="opacity:.75;border-color:#ef4444">'
                    f'<div class="obsoleto-banner">⛔ OBSOLETO</div><br>'
                    f'<span style="font-weight:700;color:#1a237e">{doc["nombre"]}</span>&nbsp;&nbsp;'
                    f'{tipo_badge(doc.get("tipo",""))}<br>'
                    f'<span style="font-size:.8rem;color:#6b7280">'
                    f'Ingresado: {doc.get("fecha_ingreso","")} &nbsp;·&nbsp; '
                    f'Archivado: {doc.get("fecha_obsolescencia","")}</span><br><br>'
                    f'<b style="font-size:.82rem">ISO 9001:</b> '
                    f'<span style="font-size:.82rem;color:#6b7280">{doc.get("iso9001","")[:80]}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                oc1, oc2 = st.columns([4,1])
                with oc2:
                    if st.button("♻️ Restaurar", key=f"rest_{doc['hash']}", use_container_width=True,
                                 help="Marcar como activo nuevamente"):
                        for i,d in enumerate(lista_maestra):
                            if d.get("hash")==doc["hash"]:
                                lista_maestra[i]["estado"] = "activo"
                                lista_maestra[i].pop("fecha_obsolescencia",None)
                                break
                        save_json(LISTA_MAESTRA_PATH, lista_maestra)
                        st.success(f"**{doc['nombre']}** restaurado como activo."); st.rerun()
                st.markdown("<hr style='margin:8px 0;border-color:#e5e7eb!important'>", unsafe_allow_html=True)


# ─── TAB: CREAR DOCUMENTO ────────────────────────────────────────────────────
def tab_crear_documento(lista_maestra: list, analisis_cache: dict):
    st.markdown('<p class="sec-title">📝 Crear / Modificar Documento SGI</p>', unsafe_allow_html=True)
    st.markdown(
        '<div class="card-info" style="color:#1e293b!important">'
        'Describí en el chat qué documento querés crear o qué mejoras aplicar. '
        'Indicá el tipo (procedimiento, instructivo, registro…), el proceso que cubre y '
        'en qué documentación existente basarte. Luego hacé clic en <b>📄 Generar Documento</b> '
        'para obtener el archivo DOCX con el formato SGI completo: '
        'encabezado, pie de página con campos de firma y toda la estructura normativa.</div>',
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)

    cl, cr = st.columns([3,5], gap="large")

    with cl:
        st.markdown('<p class="rpanel-lbl">Documentos de referencia (opcional)</p>', unsafe_allow_html=True)
        activos = [d["nombre"] for d in lista_maestra if d.get("estado","activo")=="activo"]
        base_docs = st.multiselect("Seleccioná documentos base del SGI:",
                                   options=activos, default=[],
                                   help="Se usarán como contexto para mantener coherencia de estilo")
        if st.button("🗑️ Limpiar chat", use_container_width=True):
            st.session_state["create_messages"] = []
            st.session_state["generated_doc"] = None
            st.rerun()
        st.markdown("---")
        st.markdown('<p class="rpanel-lbl">Ejemplos de solicitudes</p>', unsafe_allow_html=True)
        ejemplos = [
            "Crear un procedimiento de auditorías internas según ISO 9001 cláusula 9.2",
            "Procedimiento de control de documentos y registros para el SGI",
            "Instructivo de gestión de no conformidades y acciones correctivas",
            "Política de seguridad vial alineada con ISO 39001",
        ]
        for ej in ejemplos:
            if st.button(f"💡 {ej[:45]}...", key=f"ej_{hash(ej)}", use_container_width=True,
                         help=ej):
                st.session_state.setdefault("create_messages",[]).append({"role":"user","content":ej})
                st.rerun()

    with cr:
        if "create_messages" not in st.session_state:
            st.session_state["create_messages"] = []
        msgs = st.session_state["create_messages"]

        # Historial de chat
        chat_area = st.container(height=360)
        with chat_area:
            if not msgs:
                st.markdown(
                    '<div class="card-info" style="color:#1e293b!important">'
                    '🤖 <b>Asistente listo.</b> Describí qué documento necesitás y yo te ayudaré a '
                    'estructurarlo correctamente según las normas ISO aplicables.</div>',
                    unsafe_allow_html=True,
                )
            for msg in msgs:
                with st.chat_message(msg["role"]): st.markdown(msg["content"])

        user_input = st.chat_input("Describí el documento que querés crear o modificar...", key="create_chat")
        if user_input:
            msgs.append({"role":"user","content":user_input})
            # Respuesta de orientación antes de generar
            orient_prompt = f"""Eres experto en documentos SGI (ISO 9001:2015 / ISO 39001:2015).
El usuario quiere crear un documento: "{user_input}"
Dale una respuesta breve y orientadora (3-4 oraciones): confirmá qué tipo de documento es,
qué cláusulas ISO aplican y qué secciones principales tendrá. Luego indicale que puede hacer clic en
'Generar Documento' para obtener el archivo completo."""
            with chat_area:
                with st.chat_message("user"): st.markdown(user_input)
                with st.chat_message("assistant"):
                    with st.spinner("Analizando solicitud..."):
                        orient = _call_gemini(orient_prompt) or "Entendido. Hacé clic en **Generar Documento** para crear el archivo."
                    st.markdown(orient)
            msgs.append({"role":"assistant","content":orient})
            st.session_state["create_messages"] = msgs
            st.rerun()

        # Botón generar
        if msgs:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("📄 Generar Documento SGI (DOCX)", type="primary", use_container_width=True):
                with st.spinner("Generando contenido con Gemini y construyendo el documento DOCX..."):
                    # Contexto de documentos base
                    base_context = ""
                    for dname in base_docs:
                        texto = analisis_cache.get(dname,{}).get("_texto","") or get_document_text_from_chroma(dname)
                        if texto: base_context += f"\n\n=== {dname} ===\n{texto[:1500]}"

                    content = generate_document_content(msgs, base_context)
                    if not content:
                        st.error("No se pudo generar el contenido. Verifique la API Key de Gemini.")
                    else:
                        docx_bytes = build_docx(content)
                        st.session_state["generated_doc"] = (content, docx_bytes)

        # Mostrar resultado generado
        gen = st.session_state.get("generated_doc")
        if gen:
            content, docx_bytes = gen
            st.markdown("---")
            st.markdown(
                f'<div class="card-ok" style="color:#1e293b!important">'
                f'✅ <b>Documento generado:</b> {content.get("titulo","")}<br>'
                f'<small>Código: {content.get("codigo","")} &nbsp;·&nbsp; '
                f'Rev: {content.get("revision","")} &nbsp;·&nbsp; '
                f'Fecha: {content.get("fecha","")}</small></div>',
                unsafe_allow_html=True,
            )
            fname = content.get("titulo","documento_sgi").replace(" ","_")[:40] + ".docx"
            st.download_button(
                label="⬇️ Descargar Documento DOCX",
                data=docx_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
            st.markdown(
                '<div class="card-warn" style="color:#1e293b!important">⚠️ <b>Importante:</b> '
                'El documento se genera sin firmas. Debe circular internamente para su revisión, '
                'aprobación y firma según el procedimiento de control de documentos del SGI.</div>',
                unsafe_allow_html=True,
            )


# ─── TAB: CHATBOT ─────────────────────────────────────────────────────────────
def tab_chatbot(lista_maestra: list):
    st.markdown('<p class="sec-title">💬 Chatbot de Consulta SGI</p>', unsafe_allow_html=True)
    if not lista_maestra:
        st.markdown('<div class="card-info" style="color:#1e293b!important">👆 Suba documentos del SGI '
                    'para activar el chatbot con RAG.</div>', unsafe_allow_html=True)
        return
    chat_box = st.container(height=420)
    with chat_box:
        if not st.session_state.get("messages"):
            st.markdown(
                f'<div class="card-info" style="color:#1e293b!important">'
                f'💡 <b style="color:#1a237e;">Chatbot activo.</b> '
                f'Tengo <b>{len([d for d in lista_maestra if d.get("estado","activo")=="activo"])}</b> '
                f'documento(s) activo(s) indexado(s).<br>'
                f'Realizá una consulta sobre el Manual o las normas ISO, procedimientos, '
                f'indicadores, auditorías internas o gestión de riesgos viales.</div>',
                unsafe_allow_html=True,
            )
        for msg in st.session_state.get("messages",[]):
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
    if question := st.chat_input("Realizá una consulta sobre el Manual o las normas ISO..."):
        st.session_state.setdefault("messages",[]).append({"role":"user","content":question})
        with chat_box:
            with st.chat_message("user"): st.markdown(question)
            with st.chat_message("assistant"):
                with st.spinner("Buscando en documentos..."):
                    answer = rag_query(question)
                st.markdown(answer)
        st.session_state["messages"].append({"role":"assistant","content":answer})
    _, btn_col = st.columns([5,1])
    with btn_col:
        if st.button("🗑️ Limpiar", use_container_width=True):
            st.session_state["messages"] = []; st.rerun()


# ─── PANEL DERECHO ────────────────────────────────────────────────────────────
def render_right_panel(lista_maestra: list, analisis_cache: dict):
    st.markdown('<p class="sec-title">🔍 Panel de Análisis</p>', unsafe_allow_html=True)
    activos = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
    if not activos:
        st.markdown('<div class="card" style="text-align:center;color:#6b7280;padding:30px"><br>📌<br><br>'
                    'Los paneles aparecerán aquí<br>al subir documentos.<br><br></div>',
                    unsafe_allow_html=True)
        return
    selected = st.selectbox("Documento:", [d["nombre"] for d in activos],
                            label_visibility="collapsed")
    if not selected: return
    doc_info = next((d for d in activos if d["nombre"]==selected), {})
    analysis = analisis_cache.get(selected, {})
    st.markdown(
        f'<div class="card"><div style="font-size:.82rem;color:#6b7280;margin-bottom:4px">Seleccionado</div>'
        f'<div style="font-weight:700;color:#1a237e;font-size:.93rem;margin-bottom:8px">{selected}</div>'
        f'{tipo_badge(doc_info.get("tipo","N/A"))}</div>',
        unsafe_allow_html=True,
    )
    with st.expander("📌 Clasificación Normativa", expanded=True):
        st.markdown(f'<p class="rpanel-lbl">ISO 9001:2015</p>'
                    f'<div class="card-info" style="font-size:.84rem;color:#1e293b!important">'
                    f'{doc_info.get("iso9001","No aplica")}</div>'
                    f'<p class="rpanel-lbl">ISO 39001:2015</p>'
                    f'<div class="card-info" style="font-size:.84rem;color:#1e293b!important">'
                    f'{doc_info.get("iso39001","No aplica")}</div>',
                    unsafe_allow_html=True)
    incongruencias_doc = [t for t in analysis.get("incongruencias",[])
                          if not any(p in t for p in ["⚠️","Configure","Error"])]
    with st.expander(f"⚠️ Incongruencias ({len(incongruencias_doc)})", expanded=True):
        for i,item in enumerate(incongruencias_doc,1):
            st.markdown(f'<div class="card-warn" style="font-size:.83rem;color:#1e293b!important">'
                        f'<b>{i}.</b> {item}</div>', unsafe_allow_html=True)
        if not incongruencias_doc:
            st.info("Sin análisis disponible.")
    sugerencias_doc = [t for t in analysis.get("sugerencias",[])
                       if not any(p in t for p in ["📌","Configure","Error"])]
    with st.expander(f"📈 Sugerencias / KPIs ({len(sugerencias_doc)})", expanded=True):
        for i,item in enumerate(sugerencias_doc,1):
            st.markdown(f'<div class="card-ok" style="font-size:.83rem;color:#1e293b!important">'
                        f'<b>{i}.</b> {item}</div>', unsafe_allow_html=True)
        if not sugerencias_doc:
            st.info("Sin sugerencias disponibles.")
    st.markdown('<p style="font-size:.7rem;color:#9ca3af;text-align:center;margin-top:8px">'
                '💡 Análisis leído desde caché · Costo: <b>$0 tokens</b></p>',
                unsafe_allow_html=True)


# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    st.markdown(CSS, unsafe_allow_html=True)

    # ── Precarga del motor de embeddings (primera vez muestra spinner y recarga) ──
    if "chroma_ready" not in st.session_state:
        with st.spinner("⚙️ Iniciando motor de búsqueda semántica (~90 MB, solo la primera vez)..."):
            _init_chroma()
        st.session_state["chroma_ready"] = True
        st.rerun()   # re-renderiza la UI completa con el modelo ya en caché

    # Session state
    for k, v in [("messages",[]),("create_messages",[]),("generated_doc",None),
                 ("inc_selected",None),("last_revision",None),("last_rev_original",None),
                 ("last_sugerencia",None)]:
        if k not in st.session_state: st.session_state[k] = v

    # Cargar datos
    lista_maestra: list[dict] = load_json(LISTA_MAESTRA_PATH, [])
    # Migración: agregar campo estado a docs sin él
    changed = False
    for d in lista_maestra:
        if "estado" not in d:
            d["estado"] = "activo"; d["fecha_ingreso"] = d.get("fecha_ingreso",""); changed = True
    if changed: save_json(LISTA_MAESTRA_PATH, lista_maestra)

    analisis_cache: dict = load_json(ANALISIS_PATH, {})
    incongruencias: dict = load_incongruencias(analisis_cache)

    render_sidebar(lista_maestra, incongruencias)

    # Header con KPIs
    activos = [d for d in lista_maestra if d.get("estado","activo")=="activo"]
    n_ab = sum(1 for v in incongruencias.values() if v.get("estado")=="abierta")
    try:    n_chunks = get_collection().count()
    except: n_chunks = 0
    api_st = "✅ Activa" if os.environ.get("GEMINI_API_KEY") else "❌ Sin clave"

    st.markdown(
        f"""<div class="sgi-header">
            <h1>🏛️ Asistente de Auditoría SGI</h1>
            <p>Sistema de Gestión Integrado &nbsp;·&nbsp; ISO 9001:2015 &nbsp;·&nbsp;
               ISO 39001:2015 &nbsp;·&nbsp; RAG Local · Gemini 2.5 Flash · python-docx</p>
            <div class="kpi-row">
                <div class="kpi-badge"><span class="kv">{len(activos)}</span><span class="kl">Docs Activos</span></div>
                <div class="kpi-badge"><span class="kv">{n_ab}</span><span class="kl">Inc. Abiertas</span></div>
                <div class="kpi-badge"><span class="kv">{n_chunks}</span><span class="kl">Chunks RAG</span></div>
                <div class="kpi-badge"><span class="kv" style="font-size:.95rem">{api_st}</span><span class="kl">Gemini</span></div>
            </div>
        </div>""",
        unsafe_allow_html=True,
    )

    if not os.environ.get("GEMINI_API_KEY"):
        st.warning("⚠️ **GEMINI_API_KEY no configurada.** PowerShell: "
                   "`$env:GEMINI_API_KEY=\"tu_clave\"` → `streamlit run app.py`")

    # Panel de análisis → sidebar (evita triple anidamiento de columns)
    with st.sidebar:
        st.markdown("---")
        render_right_panel(lista_maestra, analisis_cache)

    tab1,tab2,tab3,tab4,tab5,tab6,tab7,tab8,tab9 = st.tabs([
        "📂  Documentos",
        "📋  Lista Maestra",
        "🔄  Revisiones",
        "⚠️  Incongruencias",
        "🗄️  Repositorio",
        "📅  Agenda SGI",
        "🔍  Auditoría",
        "📝  Crear Doc",
        "💬  Chatbot",
    ])
    with tab1:
        st.markdown("<br>", unsafe_allow_html=True)
        lista_maestra, analisis_cache, incongruencias = tab_documentos(
            lista_maestra, analisis_cache, incongruencias)
    with tab2:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_lista_maestra(lista_maestra)
    with tab3:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_revisiones(lista_maestra, analisis_cache)
    with tab4:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_incongruencias(lista_maestra, analisis_cache, incongruencias)
    with tab5:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_repositorio(lista_maestra, analisis_cache)
    with tab6:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_agenda_sgi(lista_maestra)
    with tab7:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_auditoria(lista_maestra, analisis_cache)
    with tab8:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_crear_documento(lista_maestra, analisis_cache)
    with tab9:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_chatbot(lista_maestra)


if __name__ == "__main__":
    main()
